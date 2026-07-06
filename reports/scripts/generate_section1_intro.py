import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = RGBColor(16, 44, 87)
    return p

def generate_section1_doc():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(30, 30, 30)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("Real-Time Embedded Remaining Useful Life Prediction for Electric Vehicle LFP Batteries via Lightweight Decision Trees and Pole-Zero Tracking")
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(16, 44, 87)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(16)
    r_sub = p_sub.add_run("Section I: Introduction Draft (Strictly ~695 words)")
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True

    add_heading_1(doc, "I. INTRODUCTION")
    
    paras = [
        "Electric vehicle adoption has surged globally as automakers race to meet emission mandates and lower battery manufacturing costs. Look inside commercial EV fleets today from major automakers like BYD, Tesla, or Ford, and you will find Lithium Iron Phosphate (LiFePO4 or LFP) battery packs dominating production lines. Automakers prefer LFP over older Nickel-Manganese-Cobalt chemistries for three reasons: cobalt is expensive, nickel cells degrade rapidly under fast charging, and LFP packs survive over 2,000 deep cycles without thermal runaway risks.",
        "However, integrating LFP cells into vehicle hardware introduces a diagnostic hurdle for embedded engineers. When a vehicle operates on the road, the onboard Battery Management System continuously samples terminal voltage, current, and temperature to track State of Health and estimate Remaining Useful Life. With older nickel packs, terminal voltage slopes steadily downward across hundreds of cycles, providing a clear tracking signal. With LFP chemistry, terminal voltage barely moves.",
        "Why does LFP voltage stay flat? Electrochemistry dictates that lithium intercalation inside an LFP lattice occurs across a two-phase equilibrium. This reaction clamps open-circuit potential to an almost horizontal plateau spanning 15% to 95% State of Charge. If you test an LFP cell at cycle 100 and check it again after 500 fast-charging sessions, the voltage profile looks practically identical. Because of this plateau, standard tracking algorithms—especially linear capacity fade formulas, Equivalent Circuit Models, and Extended Kalman Filters—frequently lose track of electrode aging. When forced to operate on flat LFP signals, these formulas fail to detect wear until the battery reaches its retirement boundary around 83% health, where capacity suddenly plunges toward failure in just two dozen cycles.",
        "To bypass flat voltage curves, many laboratories recently turned to deep learning networks like Long Short-Term Memory architectures. While neural networks achieve low prediction errors on desktop GPUs, they show serious flaws against automotive embedded constraints. Deep recurrent networks require heavy matrix multiplications taking around 45 milliseconds per inference step on standard microchips. Worse yet, neural networks function as uninterpretable black boxes lacking formal statistical uncertainty bounds. If an algorithm predicts 600 remaining cycles, a vehicle controller cannot tell whether that prediction carries a confidence bracket of +/-20 cycles or +/-300 cycles. This absence of calibrated confidence makes neural networks nearly impossible to certify under ISO 26262 automotive safety standards.",
        "Furthermore, existing literature suffers from the danger of single-point static forecasting. In early estimation studies, regression algorithms analyze the first 100 cycles of life and output a single, static health prediction at cycle 100. In an actual electric car, static predictions are dangerous. If an EV owner drives gently in year one but switches to aggressive fast-charging in year two, a static model remains blind to accelerated wear. A practical Battery Management System requires a dynamic prognostic horizon that wakes up periodically, samples recent operational wear, and updates the remaining countdown trajectory.",
        "In this research, we bridge laboratory electrochemical degradation science and embedded automotive microcontroller deployment. Rather than relying on static single-point guesses or uncalibrated neural networks, we engineer a lightweight LightGBM gradient boosting architecture making five core contributions:"
    ]

    for p_text in paras:
        doc.add_paragraph(p_text)

    bullets = [
        ("Dynamic Rolling-Window Feature Extraction at 5-Cycle Resolution: ", "We replace static predictions with a continuous tracking loop. By evaluating 8 physical variables—including internal ohmic resistance and differential capacity logarithmic variance—over a rolling 10-cycle window every 5 operating cycles, our pipeline catches capacity drops weeks before voltage shifts."),
        ("Grouped Leave-Cells-Out Validation Preventing Data Leakage: ", "To eliminate random shuffling leakage common in battery literature, we benchmark our framework across 124 commercial LFP cells (22,474 checkpoints) using strict physical cell grouping across 100 training cells and 24 unseen test cells."),
        ("Conformal Prediction Safety Brackets (+/-122 Cycles at 90% Coverage): ", "We integrate split conformal prediction over validation residuals providing vehicle firmware with a mathematical worst-case lower bound (+/-122 cycles), ensuring maintenance alerts trigger well before physical cell failure."),
        ("ECM-Based Pole-Zero Migration as a Physics Interpretability Layer: ", "We correlate empirical feature rankings directly to first-order Equivalent Circuit Model transfer functions proving mathematically that SEI film thickening shifts discrete system poles from -0.1285 rad/s toward the origin."),
        ("Sub-Millisecond Inference Feasibility on ARM Cortex Microcontrollers: ", "We conduct hardware profiling demonstrating that our compiled decision tree logic completes full inference in 0.95 ms with an 809 KB Flash footprint—executing 47x faster than deep recurrent networks within automotive budgets.")
    ]

    total_words = sum(len(p.split()) for p in paras) + sum(len((b[0]+b[1]).split()) for b in bullets)
    print(f"Total Section 1 Word Count: {total_words} words.")

    for b_title, b_desc in bullets:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.add_run(b_title).bold = True
        p_b.add_run(b_desc)

    os.makedirs(r"reports\docs", exist_ok=True)
    out_path = r"reports\docs\Section_I_Introduction_700Words.docx"
    doc.save(out_path)
    print(f"Section 1 Draft saved successfully to: {out_path}")

if __name__ == "__main__":
    generate_section1_doc()
