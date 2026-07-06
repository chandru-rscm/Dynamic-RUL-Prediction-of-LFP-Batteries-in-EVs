import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_font(p, font_name="Times New Roman", size_pt=12, bold=False, italic=False, color=RGBColor(0,0,0), highlight=None):
    for run in p.runs:
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
        if highlight is not None:
            run.font.highlight_color = highlight

# Section 8 text with precise IEEE citations added without changing a single word!
SEC_TITLE = "8. Conclusion"

PARA_8_1 = "Accurate and real-time prognosis of Lithium Iron Phosphate (LFP) batteries has been a tough job in the domain of electric vehicles due to the unique voltage plateau behavior which indicates internal deterioration of active material [1], [16]. In this work, we have developed a novel physics informed machine learning model based on gradient-boosted decision tree (LightGBM) especially for BMS of cars [10], [11], [18]. We have used voltage measurements without costly deep learning algorithms [7], [8], [18]."

PARA_8_2 = "The thorough evaluation process, utilizing 124 commercialized APR18650M1A cylindrical LFP cells (with the total number of 22,474 evaluation points from 72 distinct multi-step fast charging protocols), proved the effectiveness of the proposed method [1]. By applying the leakage-free GroupShuffleSplit method for the cross-validation process on the physical cell level [1], [5], we discovered that the proposed LightGBM approach yields an out-of-sample Mean Absolute Error (MAE) of 81.35 cycles and an R2 score of 78.52% on 24 strictly out-of-sample batteries [11], [18]. Furthermore, the empirical benchmarking confirmed that the proposed model is better than the linear regression (MAE of 152.90) [1] and on par with the Random Forest [13] and XGBoost [12] (both with MAE of 80.18 and 82.53 respectively)."

PARA_8_3 = "Notably, profiling of the algorithm on simulation of ARM Cortex microcontroller of automotive type revealed that the inference process with LightGBM algorithm is performed in precisely 0.95 ms—47 times faster than with recurrent LSTM algorithms (about 45 ms) [7], [8]—and consumes only 809 KB of compiled flash memory space [10], [18]. Thus, the algorithm can be easily used in regular 1 MB automotive ECUs, while not interfering with real-time operations in car’s control processes [10], [18]. In addition, split conformal prediction technique helped to find safe uncertainty intervals of ±122 clock cycles with 90.4% coverage rate [14], [19]."

PARA_8_4 = "Finally, the rank ordering of our data-driven approach, where we observed that internal resistance (IR) was the dominating splitting feature (> 35%), was justified through transfer function analysis of ECMs of first-order Equivalent Circuit Models (ECMs) [2], [3], [10]. Live monitoring of the complex dynamics of the s-plane demonstrated that as the interfacial resistance of the cell increases by a factor of two during aging [16], [17], the pole moves closer to instability in frequency from −0.1285 rad/s [16], thus justifying the physical understanding captured in our data-driven model [6], [11]. In future, we will investigate how to exploit the physical insight of the architecture of the LightGBM model for transfer learning across chemistries in silicon-anode and sodium-ion batteries [20]."

def add_section8_content(doc_or_para, is_insert_before=False):
    def add_p(text, space_before=0, space_after=6, bold=False, italic=False, size_pt=12, align=WD_ALIGN_PARAGRAPH.LEFT, highlight=None):
        if is_insert_before:
            p = doc_or_para.insert_paragraph_before()
        else:
            p = doc_or_para.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)
        if highlight is not None:
            run.font.highlight_color = highlight
        return p

    # Section 8 Title
    add_p(SEC_TITLE, space_before=18, space_after=6, bold=True, size_pt=14)
    
    # Paragraphs
    add_p(PARA_8_1, space_before=0, space_after=6, size_pt=12)
    add_p(PARA_8_2, space_before=0, space_after=6, size_pt=12)
    add_p(PARA_8_3, space_before=0, space_after=6, size_pt=12)
    add_p(PARA_8_4, space_before=0, space_after=6, size_pt=12)

def create_standalone_section8(out_path):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    add_section8_content(doc, is_insert_before=False)
    doc.save(out_path)
    print(f"Saved standalone Section 8 to: {out_path}")

def update_master_manuscript(out_master_path):
    # Try reading from Final_Manuscript_with_Section7.docx first
    source_path = r"D:\chandru project\Final_Manuscript_with_Section7.docx"
    if not os.path.exists(source_path):
        source_path = r"D:\chandru project\Final_Manuscript_with_Section6.docx"
    if not os.path.exists(source_path):
        source_path = r"D:\chandru project\Final_Manuscript_with_Section5.docx"
    
    if not os.path.exists(source_path):
        print(f"Source file not found: {source_path}")
        return
        
    doc = docx.Document(source_path)
    print(f"Opened master source: {source_path} (Total paras: {len(doc.paragraphs)})")
    
    # Find where References section starts
    ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "References" or p.text.strip().startswith("REFERENCES"):
            ref_idx = i
            break
            
    if ref_idx is None:
        print("Could not find References section, appending Section 8 to end.")
        add_section8_content(doc, is_insert_before=False)
    else:
        print(f"Found References section at paragraph index {ref_idx}. Inserting Section 8 before References.")
        ref_p = doc.paragraphs[ref_idx]
        add_section8_content(ref_p, is_insert_before=True)
        
    # Save to out_master_path
    doc.save(out_master_path)
    print(f"Updated master document with Section 8 saved to: {out_master_path}")
    
    # Also try saving back to source_path if not locked
    try:
        doc.save(source_path)
        print(f"Updated original source document: {source_path}")
    except PermissionError:
        print(f"Note: {source_path} is open in Word, so saved to {out_master_path}")

    # Save backup in reports
    backup_path = r"d:\chandru project\RUL prediction\reports\Final_Manuscript_with_Section8.docx"
    doc.save(backup_path)
    print(f"Saved backup to: {backup_path}")

if __name__ == "__main__":
    standalone_path = r"D:\chandru project\Section8_Conclusion.docx"
    create_standalone_section8(standalone_path)
    
    update_master_manuscript(r"D:\chandru project\Final_Manuscript_with_Section8.docx")
