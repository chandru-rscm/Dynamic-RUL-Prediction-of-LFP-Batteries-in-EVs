import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = RGBColor(16, 44, 87)
    return p

def generate_section4_doc():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(30, 30, 30)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("Section IV: Physics-Informed Feature Engineering (Review Draft)")
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(16, 44, 87)

    add_heading_1(doc, "IV. PHYSICS-INFORMED FEATURE ENGINEERING")

    p1 = ("Feature engineering for battery remaining useful life prediction bridges the gap between raw time-series sensor streams and interpretable electrochemical degradation mechanisms. Instead of feeding raw, uncalibrated voltage and current profiles directly into an algorithm—an approach that forces the machine learning model to independently discover complex thermodynamic relationships—we extract eight physically meaningful features. Each parameter directly quantifies a recognized degradation mode such as active material loss, solid electrolyte interphase thickening, or lithium plating. By embedding domain knowledge into the input feature space, we drastically reduce problem dimensionality and ensure that feature importance scores reflect physical reality.")
    doc.add_paragraph(p1)

    p2 = ("To capture dynamic battery aging without overloading onboard microcontrollers, all eight features are calculated over a rolling window of ten consecutive charge-discharge cycles. Empirical evaluations justify this ten-cycle window as the optimal operational sweet spot. Shorter windows of three to five cycles introduce excessive statistical noise from sensor fluctuations and slight ambient temperature shifts. Conversely, extended windows of twenty or fifty cycles overly smooth the localized degradation rates and mask abrupt trajectory shifts. Evaluating these features over a ten-cycle sliding window captures short-term aging acceleration while maintaining a stable, noise-resilient input stream for the decision tree engine.")
    doc.add_paragraph(p2)

    # Add Table
    p_tab_intro = doc.add_paragraph()
    r_ti = p_tab_intro.add_run("Table I: Physics-Informed Feature Set and Computation Formulas")
    r_ti.font.bold = True
    r_ti.font.size = Pt(11)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ["Feature Symbol", "Mathematical Definition", "Physical Meaning & Degradation Link", "Computation Formula"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "102C57")
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    features_data = [
        ("Q_dis", "Mean Discharge Capacity", "Tracks absolute capacity fade; primary macro-level proxy for overall State of Health.", "Mean of Q_dis(k) over window W=10"),
        ("IR", "Internal Resistance", "Measures ohmic + SEI conduction resistance; primary indicator of electrolyte decomposition.", "Mean of dV/dI during pulse at 10% SOC"),
        ("dQ_log_var", "IC Curve Log-Variance", "Captures thermodynamic peak broadening in incremental capacity curve; early sentinel.", "Log of variance of dQ/dV over window W"),
        ("Q_var", "Capacity Variance", "Quantifies cycle-to-cycle capacity stability; spikes abruptly near the aging knee point.", "Variance of Q_dis(k) over window W=10"),
        ("Q_dot", "Capacity Fade Rate", "Linear slope of capacity loss; measures local degradation speed rather than cumulative loss.", "Linear regression slope of Q_dis vs cycle k"),
        ("T_mean", "Mean Discharge Temp", "Monitors thermal stress during operational discharge; drives Arrhenius aging acceleration.", "Mean core/surface temperature over window W"),
        ("t_charge", "Time to 80% SOC", "Fast-charge duration; increases as rising internal impedance throttles constant-current phase.", "Mean charging duration to reach 80% SOC"),
        ("V_min", "End-of-Discharge Voltage", "Reflects depth of discharge polarization and rising overpotential under heavy load.", "Mean minimum terminal voltage across window")
    ]

    for row_data in features_data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)

    p3 = ("Among the eight engineered metrics, the log-variance of the incremental capacity curve serves as our most powerful early-warning sentinel, setting this framework apart from conventional State of Health tracking. Incremental capacity analysis differentiates capacity with respect to voltage (dQ/dV), converting characteristic open-circuit voltage plateaus into distinct electrochemical phase-transition peaks. As a lithium iron phosphate cell undergoes aging, active lithium loss and internal lattice strain cause these phase-transition peaks to flatten, shift, and broaden. Tracking the logarithmic variance of the dQ/dV curve identifies this structural broadening dozens of cycles before macroscopic capacity degradation becomes visible, providing critical lead time for fleet maintenance algorithms.")
    doc.add_paragraph(p3)

    p4 = ("Analyzing gradient-boosted feature gain reveals that internal resistance dominates the model's decision hierarchy, accounting for over thirty-five percent of total predictive split gain. This empirical ranking directly reinforces the physical foundation of our approach. In commercial lithium iron phosphate cells, internal resistance rises steadily as the solid electrolyte interphase thickens and consumes conductive electrolyte. The strong dominance of internal resistance in data-driven split gain directly bridges empirical machine learning with electrochemical theory, laying the exact groundwork for our equivalent circuit model pole-zero verification in subsequent sections.")
    doc.add_paragraph(p4)

    paras_text = [p1, p2, p3, p4]
    total_words = sum(len(p.split()) for p in paras_text)
    print(f"Total Section 4 Text Word Count: {total_words} words.")

    os.makedirs(r"reports\docs", exist_ok=True)
    out_path = r"reports\docs\Section_IV_Feature_Engineering.docx"
    doc.save(out_path)
    print(f"Section 4 Draft saved successfully to: {out_path}")

if __name__ == "__main__":
    generate_section4_doc()
