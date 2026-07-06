import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = RGBColor(16, 44, 87)
    return p

def generate_section2_doc():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(30, 30, 30)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("Section II: Dataset and Problem Formulation (Word Review Draft)")
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(16, 44, 87)

    add_heading_1(doc, "II. DATASET AND PROBLEM FORMULATION")
    
    # Strictly ~450 words total, zero citation numbers
    paras = [
        "To evaluate our embedded prognostic framework under realistic automotive fast-charging stress, we utilize the comprehensive open-access battery degradation dataset generated jointly by Stanford University, MIT, and the Toyota Research Institute. This experimental corpus consists of 124 commercial A123 Systems APR18650M1A cylindrical Lithium Iron Phosphate (LiFePO4) cells. Each physical cell possesses a nominal factory design capacity of 1.1 Ah and a nominal open-circuit voltage of 3.3 V. To replicate the harsh multi-stage fast-charging regimes demanded by modern electric vehicle owners, all 124 cells were continuously cycled inside precision environmental test chambers tightly regulated at a constant temperature of 30 deg C. Throughout the lifecycle testing, the cells were subjected to 72 distinct multi-step constant-current fast-charging policies ranging from standard 1C rates up to extreme 6C fast-charging profiles, followed by a uniform 4C constant-current discharge down to the lower cutoff threshold of 2.0 V.",
        "A rigorous prognostic architecture requires a mathematically unambiguous definition of cell failure. In accordance with international automotive battery standards, we define the physical End-of-Life (EOL) boundary, denoted as N_EOL, as the exact cycle index where the measured discharge capacity falls below 80% of its rated nominal capacity. For the APR18650M1A cells evaluated in this study, the retirement threshold corresponds to an absolute capacity of 0.88 Ah (State of Health <= 80%). Consequently, for any active operating cycle index t during the lifespan of a battery pack, the true Remaining Useful Life, RUL(t), represents the exact number of operational cycles remaining until the physical cell crosses the retirement boundary. Mathematically, this prognostic relationship is expressed as RUL(t) = N_EOL - t.",
        "Examining the empirical degradation trajectories across all 124 cells reveals fundamental electrochemical aging behaviors that govern our diagnostic design. During the early and middle stages of battery life (from cycle 1 down to approximately 88% State of Health), capacity loss progresses along an exceptionally gentle, linear slope caused by stable solid electrolyte interphase (SEI) film thickening. However, as cells approach late-life aging, the degradation trajectories exhibit a dramatic non-linear inflection, commonly referred to in battery physics as the aging knee point. At this juncture, internal lithium plating and rapid pore clogging trigger an accelerated capacity plunge, causing cells to drop from healthy operational states to complete failure in fewer than 30 cycles.",
        "Therefore, the core real-time embedded prediction task is formally stated as follows: At any diagnostic evaluation checkpoint cycle t, given an 8-dimensional physical feature vector extracted over a rolling historical window of the past 10 cycles, the onboard machine learning model must immediately estimate the remaining cycle countdown RUL(t). Furthermore, the algorithm must execute within strict microsecond execution budgets on automotive microchips while providing formal mathematical confidence brackets to alert vehicle controllers before the knee point occurs."
    ]

    for p_text in paras:
        doc.add_paragraph(p_text)

    total_words = sum(len(p.split()) for p in paras)
    print(f"Total Section 2 Word Count: {total_words} words.")

    os.makedirs(r"reports\docs", exist_ok=True)
    out_path = r"reports\docs\Section_II_Dataset_Problem_Formulation.docx"
    doc.save(out_path)
    print(f"Section 2 Draft saved successfully to: {out_path}")

if __name__ == "__main__":
    generate_section2_doc()
