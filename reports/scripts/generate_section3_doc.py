import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = RGBColor(16, 44, 87)
    return p

def generate_section3_doc():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(30, 30, 30)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("Section III: Related Work (Plain Text Word Review Draft)")
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(16, 44, 87)

    add_heading_1(doc, "III. RELATED WORK")
    
    paras = [
        "Battery Remaining Useful Life prediction has historically progressed across three distinct methodological paradigms. The earliest generation relied heavily on physics-based formulations and empirical degradation laws. Researchers utilized Arrhenius-type equations to model capacity loss under constant temperature and aging conditions. While these empirical formulas are computationally inexpensive and simple to implement in legacy firmware, they demand extensive parameter tuning for every specific cell formulation and fail to adapt to the highly dynamic, non-linear interactions of rapid fast-charging and varying thermal profiles. Conversely, detailed electrochemical models such as the Doyle-Fuller-Newman framework offer high theoretical fidelity by simulating ion transport across porous electrodes. However, solving partial differential equations in real time exceeds the processing budget of automotive microcontrollers, and these frameworks depend on internal physical quantities—such as solid diffusion coefficients and interphase film growth rates—that cannot be measured non-invasively inside an operating vehicle.",
        "The second generation shifted toward classical machine learning techniques applied to empirical battery data. Researchers explored Gaussian process regression to extract probabilistic remaining life predictions from capacity curves, though the cubic computational complexity of matrix inversion limits its real-time scalability. Support vector machines and random forests were subsequently applied to differential voltage and incremental capacity curves gathered during low-current charging, establishing that electrochemically motivated feature transformations significantly outperform raw voltage tracking. Notably, the landmark study by Severson and colleagues demonstrated that an elastic-net linear regression model, trained on variance statistics of incremental capacity curves during the initial hundred cycles, could forecast lifetime across commercial lithium iron phosphate cells with strong baseline accuracy.",
        "More recently, the literature has been dominated by deep learning architectures. Long Short-Term Memory networks and recurrent architectures have been widely implemented to process sequential voltage and current time series. While these recurrent networks achieve impressive accuracy within uniform dataset splits, they frequently overfit to specific training trajectories and struggle to generalize across diverse multi-stage charging protocols. Modern transformer models with self-attention mechanisms capture long-range temporal dependencies but introduce millions of trainable weights, rendering them entirely unfeasible for resource-constrained automotive microchips. Furthermore, deep learning frameworks operate as uninterpretable black boxes and typically lack formal statistical uncertainty quantification. When prediction intervals are reported, they commonly rely on heuristic techniques like Monte Carlo dropout, which tend to severely underestimate uncertainty boundaries when encountering out-of-distribution battery wear.",
        "To provide rigorous safety guarantees without heavy computational overhead, conformal prediction has emerged as a powerful mathematical alternative to Bayesian regression. Split conformal prediction wraps pre-trained point estimators with distribution-free prediction intervals that carry exact finite-sample coverage guarantees. Regardless of the underlying error distribution, conformal intervals ensure that true battery remaining useful life falls within predicted brackets at a user-specified confidence level. While conformal prediction is gaining traction across statistical learning domains, its application to real-time battery diagnostics remains sparse, and existing studies have not bridged conformal safety intervals with physical circuit model verification.",
        "To synthesize high predictive accuracy, computational efficiency, and physical interpretability, gradient boosted decision trees—specifically LightGBM—offer a compelling solution for tabular sensor data. LightGBM utilizes a leaf-wise tree growth strategy combined with histogram-based feature splitting and gradient-based one-side sampling. This architecture trains orders of magnitude faster than deep recurrent networks while preventing overfitting on moderate-sized battery datasets. Moreover, its native split-gain feature rankings provide transparent insights into which physical variables drive aging forecasts. To our knowledge, the framework presented in this research is the first to unite lightweight LightGBM decision trees, rolling-window electrochemical features, grouped leakage-free validation, split-conformal safety intervals, and equivalent circuit pole-zero interpretability into a unified, embedded-ready prognostic system."
    ]

    for p_text in paras:
        doc.add_paragraph(p_text)

    total_words = sum(len(p.split()) for p in paras)
    print(f"Total Section 3 Word Count: {total_words} words.")

    os.makedirs(r"reports\docs", exist_ok=True)
    out_path = r"reports\docs\Section_III_Related_Work.docx"
    doc.save(out_path)
    print(f"Section 3 Draft saved successfully to: {out_path}")

if __name__ == "__main__":
    generate_section3_doc()
