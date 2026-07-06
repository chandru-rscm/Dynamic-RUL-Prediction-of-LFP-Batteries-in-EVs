import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX

def set_font(p, font_name="Times New Roman", size_pt=12, bold=False, italic=False, color=RGBColor(0,0,0), highlight=None):
    for run in p.runs:
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
        if highlight is not None:
            run.font.highlight_color = highlight

# Section 4 text with precise IEEE citations added without changing a single word!
SEC_TITLE = "4. Physics-Informed Feature Engineering"

SUB_4_1 = "4.1 Domain-Guided Feature Selection"
PARA_4_1 = "Feature engineering process to predict remaining useful life of batteries fills the gap between the raw sensor time series data and degradation mechanisms in terms of electrochemistry [6], [11]. Instead of applying uncalibrated values of voltage and current data, which requires the model to figure out the complex thermodynamic process on its own [10], [18], we come up with eight features related to the known degradation mechanisms such as loss of active materials, SEI formation, or lithium plating [16], [17]. Thus, we introduce our knowledge into the set of features and reduce the dimensionality of the problem [6], [11]."
TABLE_1_TEXT = "Table 1 Physics-Informed Feature Set and Computation Formulas"

SUB_4_2 = "4.2 Rolling-Window Formulation and Interval Analysis"
PARA_4_2 = "For the purpose of modeling the evolution of battery aging without putting a strain on the onboard microcontrollers [10], [18], calculations for all eight characteristics will be carried out within a moving window of ten consecutive charge-discharge cycles [1], [6]. Empirical research proves the fact that a window of ten cycles is the optimal operating range [1], [5]. Windows which contain less than ten cycles, i.e. windows from three to five cycles, are influenced by too much statistical noise due to the variability of sensors and temperature changes [10], [18]. Meanwhile, windows containing more than ten cycles, for example twenty or fifty cycles, oversimplify the rate of degradation locally [1], [5]."
TABLE_2_TEXT = "Table 2 Rolling Lookback Window (W ) Sensitivity Analysis"

SUB_4_3 = "4.3 Incremental Capacity Log-Variance as an Early Sentinel"
PARA_4_3 = "Among these eight engineering indicators, the logarithmic variance of the incremental capacity profile serves as the most dependable early indicator, which makes the proposed system distinct from conventional State of Health monitoring techniques [1], [9]. The incremental capacity measurement uses differential capacity with respect to voltage (dQ/dV) [9], [11]. This way, the open circuit voltage plateaus are converted to electrochemical phase transition peaks [1], [9]. During the lifespan of a lithium iron phosphate battery, the depletion of active lithium content and the lattice strain result in decrease of sharpness of the phase transition peaks [16], [17]. Monitoring the logarithmic variance of the dQ/dV profile allows identifying the broadening tens of cycles before any changes in capacity can be observed [1], [11]."

SUB_4_4 = "4.4 Feature Gain Analysis and Electrochemical Link"
PARA_4_4 = "As revealed by the gradient-boosted feature gain analysis, the internal resistance (IR) is the predominating variable of the set of features for constructing our decision trees with almost more than thirty-five percent of total split gains, which is quite evident from Fig. 1 [10], [18]. Such an empirical ranking gives compelling numerical evidence validating the physical theory of our model [6], [16]. For commercial lithium iron phosphate batteries, the internal resistance increases progressively due to the formation and growth of the solid electrolyte interface (SEI) layer and gradual depletion of the conducting electrolyte [16], [17]. Importantly, such feature-based dominance of IR gives a direct validation and prediction to our Equivalent Circuit Model (ECM) physics verification in Section VI [2], [3]. As shown in later migration analysis of first-order pole-zero pairs, the cell ohmic resistance (R0) doubles from the healthy to the aged states, validating the fact that the LightGBM algorithm has independently identified the major physical phenomenon causing battery end-of-life [10], [16]."
FIG_1_TEXT = "Fig. 1 LightGBM Feature Importance ranking by split count across the eight physics-informed parameters. Internal resistance (IR) clearly dominates the tree decision hierarchy, providing empirical justification for the Equivalent Circuit Model (ECM) physics verification investigated in Section VI."

def add_section4_content(doc_or_para, is_insert_before=False):
    # Helper to add paragraph either by appending (doc) or inserting before (para)
    def add_p(text, space_before=0, space_after=6, bold=False, italic=False, size_pt=12, align=WD_ALIGN_PARAGRAPH.LEFT, highlight=None):
        if is_insert_before:
            p = doc_or_para.insert_paragraph_before()
        else:
            p = doc_or_para.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)
        if highlight is not None:
            run.font.highlight_color = highlight
        return p

    # Section 4 Title
    add_p(SEC_TITLE, space_before=18, space_after=6, bold=True, size_pt=14)
    
    # 4.1
    add_p(SUB_4_1, space_before=12, space_after=4, bold=True, size_pt=12)
    add_p(PARA_4_1, space_before=0, space_after=6, size_pt=12)
    add_p(TABLE_1_TEXT, space_before=6, space_after=12, italic=True, size_pt=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, highlight=WD_COLOR_INDEX.YELLOW)
    
    # 4.2
    add_p(SUB_4_2, space_before=12, space_after=4, bold=True, size_pt=12)
    add_p(PARA_4_2, space_before=0, space_after=6, size_pt=12)
    add_p(TABLE_2_TEXT, space_before=6, space_after=12, italic=True, size_pt=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, highlight=WD_COLOR_INDEX.YELLOW)
    
    # 4.3
    add_p(SUB_4_3, space_before=12, space_after=4, bold=True, size_pt=12)
    add_p(PARA_4_3, space_before=0, space_after=6, size_pt=12)
    
    # 4.4
    add_p(SUB_4_4, space_before=12, space_after=4, bold=True, size_pt=12)
    add_p(PARA_4_4, space_before=0, space_after=6, size_pt=12)
    add_p(FIG_1_TEXT, space_before=6, space_after=12, italic=True, size_pt=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, highlight=WD_COLOR_INDEX.YELLOW)

def create_standalone_section4(out_path):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    add_section4_content(doc, is_insert_before=False)
    doc.save(out_path)
    print(f"Saved standalone Section 4 to: {out_path}")

def update_master_manuscript(out_master_path):
    # Try reading from Final_Manuscript_with_Section3.docx first
    source_path = r"D:\chandru project\Final_Manuscript_with_Section3.docx"
    if not os.path.exists(source_path):
        source_path = r"D:\chandru project\Final_Manuscript_Abstract_Intro_RelatedWork.docx"
    
    if not os.path.exists(source_path):
        print(f"Source file not found: {source_path}")
        return
        
    doc = docx.Document(source_path)
    print(f"Opened master source: {source_path} (Total paras: {len(doc.paragraphs)})")
    
    # Find where References section starts
    ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "References" or p.text.strip().startswith("REFERENCES"):
            ref_idx = i
            break
            
    if ref_idx is None:
        print("Could not find References section, appending Section 4 to end.")
        add_section4_content(doc, is_insert_before=False)
    else:
        print(f"Found References section at paragraph index {ref_idx}. Inserting Section 4 before References.")
        ref_p = doc.paragraphs[ref_idx]
        add_section4_content(ref_p, is_insert_before=True)
        
    # Save to out_master_path
    doc.save(out_master_path)
    print(f"Updated master document with Section 4 saved to: {out_master_path}")
    
    # Also try saving back to source_path if not locked
    try:
        doc.save(source_path)
        print(f"Updated original source document: {source_path}")
    except PermissionError:
        print(f"Note: {source_path} is open in Word, so saved to {out_master_path}")

    # Save backup in reports
    backup_path = r"d:\chandru project\RUL prediction\reports\Final_Manuscript_with_Section4.docx"
    doc.save(backup_path)
    print(f"Saved backup to: {backup_path}")

if __name__ == "__main__":
    standalone_path = r"D:\chandru project\Section4_Physics_Informed_Feature_Engineering.docx"
    create_standalone_section4(standalone_path)
    
    update_master_manuscript(r"D:\chandru project\Final_Manuscript_with_Section4.docx")
