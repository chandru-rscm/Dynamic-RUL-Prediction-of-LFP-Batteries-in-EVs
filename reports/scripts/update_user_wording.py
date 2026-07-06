import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def update_doc_exact_user_wording():
    path1 = r"D:\chandru project\Final_Manuscript_Abstract_Intro_RelatedWork.docx"
    path2 = r"reports\Final_Manuscript_Abstract_Intro_RelatedWork.docx"
    
    doc = docx.Document(path1)
    
    # Extract paragraphs starting from Related Work (3 Related Work or Section 3)
    paras_from_related = []
    start_collecting = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("3 Related Work") or text.startswith("3.") and "Related Work" in text:
            start_collecting = True
        if start_collecting:
            paras_from_related.append((text, p.style.name, p.paragraph_format.left_indent))
            
    new_doc = docx.Document()
    
    # Page setup
    for section in new_doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    normal_style = new_doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)
    
    # Title
    title_p = new_doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Dynamic Remaining Useful Life (RUL) Prediction of Lithium-Iron-Phosphate Batteries in Electric Vehicles: A Lightweight Edge-Deployable Framework")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x00, 0x2B, 0x49)
    title_p.paragraph_format.space_after = Pt(14)
    
    # Authors
    auth_p = new_doc.add_paragraph()
    auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth_run = auth_p.add_run("Automotive Battery Management System (BMS) Telemetry & Edge Intelligence Laboratory")
    auth_run.font.name = 'Times New Roman'
    auth_run.font.size = Pt(11)
    auth_run.font.italic = True
    auth_p.paragraph_format.space_after = Pt(18)
    
    # Abstract Heading
    abs_h = new_doc.add_paragraph()
    abs_hr = abs_h.add_run("ABSTRACT")
    abs_hr.font.name = 'Times New Roman'
    abs_hr.font.size = Pt(12)
    abs_hr.font.bold = True
    abs_hr.font.color.rgb = RGBColor(0x00, 0x2B, 0x49)
    abs_h.paragraph_format.space_after = Pt(4)
    
    # User's exact Abstract with citations
    abs_p = new_doc.add_paragraph()
    abs_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abs_text = (
        "Accurate prediction of the Remaining Useful Life (RUL) in LiFePO4, or LFP battery chemistry, is a safety-related necessity in today’s electric vehicles [1], [5]. "
        "Whereas LFP cells exhibit better stability than any other chemistry based on nickel, the inherent flatness of the open circuit voltage curve makes diagnosis "
        "with embedded Battery Management Systems (BMS) exceptionally challenging [3], [16]. Although recent deep learning methods have reached very good accuracy, the "
        "problematic aspect is that they require enormous computational power and are not equipped with calibrated uncertainty intervals for microcontrollers [7], [8], [15]. "
        "In this paper, we propose a low-power physics-inspired machine learning algorithm based on gradient-boosted decision trees (LightGBM) [10], [11]. To predict capacity "
        "reduction before voltage decrease, the algorithm evaluates eight rolling window electrochemical parameters every five cycles [6], [16]. We validated the method on 124 "
        "commercial LFP cells subjected to harsh multi-stage fast charging protocols (22,474 checkpoints) [1], [4], and obtained extremely fast inference times (0.95 ms on "
        "ARM Cortex microcontrollers) [10], [18], as well as reliable split conformal safety prediction brackets (±122 cycles at 90% confidence) [14], [19]. "
        "Keywords: Lithium-Iron-Phosphate (LFP); Remaining Useful Life (RUL); Battery Management System (BMS); LightGBM; Conformal Prediction; Microcontrollers"
    )
    abs_p.add_run(abs_text)
    abs_p.paragraph_format.space_after = Pt(16)
    
    # Section 1: Introduction Heading
    h1 = new_doc.add_paragraph()
    h1r = h1.add_run("1. Introduction")
    h1r.font.name = 'Times New Roman'
    h1r.font.size = Pt(14)
    h1r.font.bold = True
    h1r.font.color.rgb = RGBColor(0x00, 0x2B, 0x49)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    # User's exact Introduction paragraphs with citations
    intro_paras = [
        "The use of electric vehicles has risen sharply across the world due to efforts by car manufacturers to conform to emissions regulations and reduce the cost of production of batteries [1]. When you open up the commercial fleet of electric vehicles currently produced by companies such as BYD, Tesla, or Ford, you will see that the most dominant type of batteries is the Lithium Iron Phosphate (LiFePO4 or LFP) battery pack [4], [5]. There are three reasons why car makers favor LFP batteries over Nickel-Manganese-Cobalt (NMC) batteries: cobalt is expensive and ethically controversial, nickel cells degrade rapidly under daily fast charging, and LFP packs routinely survive over 2,000 deep charge-discharge cycles without triggering thermal runaway risks [4], [5].",
        
        "But when LFP chemistry is applied in automotive equipment, this leads to a major problem of diagnostic difficulties with embedded controllers of Battery Management Systems (BMS) [2], [3]. In NMC cells, the voltage decreases drastically with the reduction of internal electrode capacity after several hundred charge/discharge cycles [13]. The decreasing trend of voltage allows onboard systems to diagnose State of Health (SOH) and Remaining Useful Life (RUL) with the help of an unambiguous tracking algorithm [5], [13]. With LFP, there is no diagnostic voltage signal at all [3], [16].",
        
        "What makes conventional voltage-based estimation inaccurate on LFP batteries? According to electrochemistry, the lithium insertion process in the LFP crystal lattice is a two-phase thermodynamic equilibrium process (LiFePO4 ↔ FePO4), which dictates a very stiff plateau of the cell open circuit voltage (3.29 V to 3.33 V) spanning around 15% to 95% State of Charge (SOC) [3], [16]. Since this voltage plateau is very stiff even when there is substantial loss of active material in the cell, estimation based on voltage measurements suffers from zero observability (∂V/∂SOH ≈ 0) [17]. Therefore, conventional BMS estimators have no ability to see any internal material degradation in the electrode until the cell reaches its critical threshold (∼ 80% health) [16], [17]. After crossing this threshold, the cell falls off the voltage plateau and goes into capacity failure in only two dozen operating cycles.",
        
        "To overcome the limitation of flat voltage curves, researchers in many labs have been using deep learning approaches such as Long Short-Term Memory (LSTM) architectures [7]–[9]. However, even though neural networks are able to predict very accurately on desktops via GPUs, they are not suited for automotive embedded systems since they suffer from major drawbacks [10], [20]. Deep recurrent networks rely on heavy matrix operations, which take roughly 45 ms per operation on standard microprocessors [2], [10]. Moreover, neural networks operate like black boxes that do not give any information regarding uncertainty measures [14], [15]. When the algorithm predicts that 600 cycles are left before cell death, it is hard to determine whether the prediction range is ±20 cycles or ±300 cycles [14], [15].",
        
        "Moreover, the existing body of literature is susceptible to the problem of single-point static forecasting [1], [5]. When estimating health during the initial phase, regression models take into account the first 100 cycles of the vehicle’s life and produce a single, static forecast at the end of the 100th cycle [1], [4]. In reality, when used in an electric car, a static model can be dangerous [5], [9]. For instance, when an EV owner charges gently in the first year of the vehicle’s life but then shifts to aggressive charging in the second year, a static model fails to recognize faster wear and tear [4], [5].",
        
        "In this work, we connect the dots between electrochemical degradation of batteries in the lab and real-world applications of microcontrollers inside cars [10], [11]. Instead of making assumptions and using untrained neural networks, we develop a compact LightGBM gradient boosting model that makes five essential contributions:\n\n"
        "• Dynamic Rolling-Window Feature Extraction at 5-Cycle Resolution: Our framework is built with continuous tracking in mind [6]. By analyzing 8 parameters including ohmic resistance and logarithm of differential capacity variance in a 10-cycle rolling window every 5 cycles, we detect capacity loss weeks ahead of voltage changes [6], [16].\n"
        "• Grouped Leave-Cells-Out Validation Without Data Leakage: To get rid of the leakage effect from random data shuffling in most battery research, we validate our framework on 124 commercial LFP cells (22,474 checkpoints) across 100 training cells and 24 testing cells [1], [5].\n"
        "• Confidence Intervals for Predicted Cycle Count Using Conformal Prediction Safety Brackets (±122 Cycles at 90% Coverage): This approach allows us to apply split conformal prediction using validation residuals, and as a result, we provide the mathematical worst-case lower bound (±122 cycles), which will guarantee that maintenance will be triggered long before the physical cell fails [14], [19].\n"
        "• Pole-Zero Migration Through Equivalent Circuit Models Physics Layer Interpretability: The empirical ranking of features is mapped through first-order Equivalent Circuit Model transfer functions; hence, mathematically proven, thickening of SEI film will change discrete system poles from −0.1285 rad/s towards the origin [2], [3].\n"
        "• Inference Time Less Than Sub-Millisecond Using ARM Cortex Microcontrollers: Hardware profiling of our compiled decision tree model shows that inference is performed in 0.95 ms while using 809 KB of Flash memory [10], [18]."
    ]
    
    for ip in intro_paras:
        p = new_doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(ip)
        
    # Append exact Related Work and References untouched
    for text, style_name, indent in paras_from_related:
        p = new_doc.add_paragraph()
        if indent is not None:
            p.paragraph_format.left_indent = indent
            if text.strip().startswith("["):
                p.paragraph_format.first_line_indent = Inches(-0.3)
                
        if text.startswith("3.") or "Related Work" in text or "Positioning of the Proposed" in text or "References" in text:
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14) if ("Related Work" in text or "References" in text) else Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x2B, 0x49)
            p.paragraph_format.space_before = Pt(12) if ("Related Work" in text or "References" in text) else Pt(8)
            p.paragraph_format.space_after = Pt(4)
        else:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            
    new_doc.save(path1)
    new_doc.save(path2)
    print("Updated document with user's exact wording + citations successfully!")

if __name__ == "__main__":
    update_doc_exact_user_wording()
