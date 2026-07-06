import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_font(p, font_name="Times New Roman", size_pt=12, bold=False, italic=False, color=RGBColor(0,0,0)):
    for run in p.runs:
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color

# Section 3 text with precise IEEE citations added without changing a single word!
PARA_1 = "For testing the performance of the proposed embedded prognostic algorithm under real-world conditions of fast charging of vehicles, the test data set that consists of the extensive battery degradation data is employed [1]. The experimental data is obtained through cooperation between Stanford University, MIT, and the Toyota Research Institute [1], [4]. This test data set contains 124 commercial A123 Systems APR18650M1A cylindrical Lithium Iron Phosphate (LiFePO4) cells [1]. Individual cell’s nominal capacity is 1.1 Ah and its nominal voltage is 3.3 V. In order to test the fast-charging protocols of harsh environment that are typical for contemporary consumers of electric cars, all 124 cells are being constantly charged in the environmental chamber, which temperature was constant (30°C) [1], [4]. During the lifetime experiment, each cell has experienced 72 various multi-stage fast-charging protocols ranging from nominal 1C rate up to 6C fast-charging protocol followed by 4C discharging down to 2.0 V [1]."

PARA_2 = "In order to develop an accurate model for prognostics, there must be a clearly defined failure point of the cell [2], [5]. In accordance with the automotive battery standard, NEOL is described as the mathematical point of the End-of-Life (EOL) of the cell, which means the specific number of cycles at which the discharge capacity falls below 80% of its nominal capacity of the cell [3], [10]. In the case of the APR18650M1A cells that have been used in this study, the retirement point is reached when the absolute capacity of the cell reaches 0.88 Ah (State of Health ≤ 80%) [1], [18]. Thus, the remaining useful life of any battery pack will be mathematically determined as follows:"

EQ_1 = "RUL(t) = NEOL – t. (1)"

PARA_3 = "By analyzing the empirical degradation paths of all 124 cells, we get a good understanding of the fundamental electrochemical processes of cell degradation, and this knowledge will be very important in developing our diagnostic system [6], [16]. At the start and at the intermediate stage of the life of the batteries (from cycle 1 to 88% State of Health), the capacity fade occurs on the very smooth linear path due to the formation of the Solid Electrolyte Interphase (SEI) film on the electrodes [16], [17]. However, at the late stage of cell aging, the paths of cells degradation drastically change their direction in what is called “aging knee point” in battery physics [1], [17]. In this stage, lithium plating and pore clogging cause rapid degradation of the cells, leading to cell failure in less than 30 cycles [16], [17]."

PARA_4 = "Herein, the basic real-time embedded prediction problem is unambiguously specified as follows: During each diagnostic assessment cycle t, based on an 8-dimensional feature vector derived from physical attributes which has already been computed by taking into account a sliding window of the past 10 cycles [6], [11], the onboard machine learning algorithm must quickly predict the number of remaining cycles until RUL(t) [10], [18]. Also, the algorithm needs to operate in very tight microsecond timing windows for automotive microchips [10], [18] and provide confidence intervals [14], [19]."

def create_standalone_section3(out_path):
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Section Title
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run("3. Dataset and Problem Formulation")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Paragraph 1
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(6)
    p1.paragraph_format.line_spacing = 1.15
    p1.add_run(PARA_1)
    set_font(p1)
    
    # Paragraph 2
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.line_spacing = 1.15
    p2.add_run(PARA_2)
    set_font(p2)
    
    # Equation 1
    eq = doc.add_paragraph()
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq.paragraph_format.space_before = Pt(6)
    eq.paragraph_format.space_after = Pt(6)
    eq_run = eq.add_run(EQ_1)
    eq_run.font.name = "Times New Roman"
    eq_run.font.size = Pt(12)
    eq_run.font.bold = True
    
    # Paragraph 3
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(6)
    p3.paragraph_format.line_spacing = 1.15
    p3.add_run(PARA_3)
    set_font(p3)
    
    # Paragraph 4
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(6)
    p4.paragraph_format.line_spacing = 1.15
    p4.add_run(PARA_4)
    set_font(p4)
    
    doc.save(out_path)
    print(f"Saved standalone Section 3 with citations to: {out_path}")

def update_master_manuscript(master_path):
    # Try opening from Final_Manuscript_with_Section3.docx if it exists, else from Abstract_Intro_RelatedWork
    source_path = r"D:\chandru project\Final_Manuscript_Abstract_Intro_RelatedWork.docx"
    if not os.path.exists(source_path):
        print(f"Source file not found: {source_path}")
        return
        
    doc = docx.Document(source_path)
    
    # Find where References section starts
    ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "References" or p.text.strip().startswith("REFERENCES"):
            ref_idx = i
            break
            
    if ref_idx is None:
        print("Could not find References section in master document, appending to end.")
        ref_idx = len(doc.paragraphs)
    else:
        print(f"Found References section at paragraph index {ref_idx}.")
        
    ref_p = doc.paragraphs[ref_idx]
    
    # Insert Section Title
    h = ref_p.insert_paragraph_before()
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run("3. Dataset and Problem Formulation")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Paragraph 1
    p1 = ref_p.insert_paragraph_before()
    p1.paragraph_format.space_after = Pt(6)
    p1.paragraph_format.line_spacing = 1.15
    p1.add_run(PARA_1)
    set_font(p1)
    
    # Paragraph 2
    p2 = ref_p.insert_paragraph_before()
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.line_spacing = 1.15
    p2.add_run(PARA_2)
    set_font(p2)
    
    # Equation 1
    eq = ref_p.insert_paragraph_before()
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq.paragraph_format.space_before = Pt(6)
    eq.paragraph_format.space_after = Pt(6)
    eq_run = eq.add_run(EQ_1)
    eq_run.font.name = "Times New Roman"
    eq_run.font.size = Pt(12)
    eq_run.font.bold = True
    
    # Paragraph 3
    p3 = ref_p.insert_paragraph_before()
    p3.paragraph_format.space_after = Pt(6)
    p3.paragraph_format.line_spacing = 1.15
    p3.add_run(PARA_3)
    set_font(p3)
    
    # Paragraph 4
    p4 = ref_p.insert_paragraph_before()
    p4.paragraph_format.space_after = Pt(6)
    p4.paragraph_format.line_spacing = 1.15
    p4.add_run(PARA_4)
    set_font(p4)
    
    # Save updated master to Final_Manuscript_with_Section3.docx
    new_master_path = r"D:\chandru project\Final_Manuscript_with_Section3.docx"
    doc.save(new_master_path)
    print(f"Updated master document with citations saved to: {new_master_path}")
    
    # Also try saving to original if not locked
    try:
        doc.save(source_path)
        print(f"Updated original master document: {source_path}")
    except PermissionError:
        print(f"Note: {source_path} is open in Word, so saved to {new_master_path}")

    # Also save backup in reports
    backup_path = r"d:\chandru project\RUL prediction\reports\Final_Manuscript_with_Section3.docx"
    doc.save(backup_path)
    print(f"Saved backup to: {backup_path}")

if __name__ == "__main__":
    standalone_path = r"D:\chandru project\Section3_Dataset_and_Problem_Formulation.docx"
    create_standalone_section3(standalone_path)
    
    update_master_manuscript(r"D:\chandru project\Final_Manuscript_with_Section3.docx")
