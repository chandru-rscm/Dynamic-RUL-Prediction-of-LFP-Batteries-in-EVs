import os
import shutil
import zipfile
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def format_table(table, col_widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if col_widths and j < len(col_widths):
                cell.width = col_widths[j]
            if i == 0:
                set_cell_background(cell, "102C57") # Navy Header
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.font.name = "Calibri"
                        r.font.size = Pt(10)
            else:
                bg = "F5F7FA" if i % 2 == 1 else "FFFFFF"
                set_cell_background(cell, bg)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(33, 37, 41)
                        r.font.name = "Calibri"
                        r.font.size = Pt(9.5)

def create_figures_zip():
    print("=== 1. CREATING FIGURES ZIP FOLDER ===")
    figures_mapping = [
        ("01_feature_importance.png", r"d:\chandru project\RUL prediction\results\figures\02_feature_importance.png"),
        ("02_flow_architecture.png", r"d:\chandru project\RUL prediction\reports\figures\flow_architecture.png"),
        ("03_polling_blind_spot.png", r"d:\chandru project\RUL prediction\reports\figures\polling_blind_spot.png"),
        ("04_pole_zero_migration.png", r"d:\chandru project\RUL prediction\reports\figures\pole_zero_migration.png"),
        ("05_true_vs_predicted_rul.png", r"d:\chandru project\RUL prediction\results\figures\01_true_vs_predicted_rul.png"),
        ("06_empirical_benchmarking.png", r"d:\chandru project\RUL prediction\results\benchmarks\model_comparison_bar_chart.png"),
        ("07_prediction_errors_histogram.png", r"d:\chandru project\RUL prediction\results\figures\03_prediction_errors_histogram.png"),
        ("08_confusion_matrix.png", r"d:\chandru project\RUL prediction\results\figures\06_confusion_matrix.png")
    ]
    
    zip_paths = [
        r"D:\chandru downloads\LFP_RUL_Manuscript_Figures.zip",
        r"d:\chandru project\RUL prediction\LFP_RUL_Manuscript_Figures.zip"
    ]
    
    for zip_path in zip_paths:
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for arcname, src_path in figures_mapping:
                if os.path.exists(src_path):
                    zf.write(src_path, arcname)
                    print(f" Added to zip: {arcname} (from {os.path.basename(src_path)})")
                else:
                    print(f" ERROR: Missing source image {src_path}")
        print(f"Saved Zip Archive to: {zip_path}\n")

def create_all_tables_doc():
    print("=== 2. CREATING STANDALONE TABLES WORD DOCUMENT ===")
    tables_data = [
        ("Table 1 Physics-Informed Feature Set and Computation Formulas", [
            ["Symbol", "Mathematical Definition", "Physical Meaning & Degradation Link", "Computation Formula"],
            ["k", "Cycle Number", "Tracks baseline temporal aging progression and cumulative operational charge-discharge duration of the cell.", "Current operational cycle index k at checkpoint"],
            ["SOH", "State of Health", "Normalized remaining capacity ratio; primary macro-level health indicator across cell form factors.", "Q_dis(k) / Q_nominal"],
            ["IR", "Internal Resistance", "Measures ohmic + SEI conduction resistance; primary indicator of electrolyte decomposition and interfacial thickening.", "Mean(dV / dI) during pulse at 10% SOC"],
            ["T_mean", "Mean Discharge Temp", "Monitors thermal stress during operational discharge; drives Arrhenius aging acceleration and SEI growth.", "Mean surface/core temperature over W = 10"],
            ["ΔSOH", "Capacity Fade Rate", "Normalized rate of capacity loss over lookback window; measures local degradation velocity (ΔSOH).", "[Q_dis(k-10) - Q_dis(k)] / Q_nominal"],
            ["ΔQ_log_var", "IC Curve Log-Variance", "Captures thermodynamic peak broadening in incremental capacity curve; acts as an early warning sentinel.", "log[Var(dQ / dV)] over window W=10"],
            ["ΔQ_min", "IC Curve Min Shift", "Maximum localized downward shift (valley) in the differential capacity profile over the rolling lookback window.", "min[Δ(dQ / dV)] over window W=10"],
            ["ΔQ_mean", "IC Curve Mean Shift", "Average baseline shift of incremental capacity curve; tracks global thermodynamic drift and stoichiometric imbalance.", "mean[Δ(dQ / dV)] over window W=10"]
        ], [Inches(0.9), Inches(1.8), Inches(2.5), Inches(1.8)]),
        
        ("Table 2 Rolling Lookback Window (W) Sensitivity Analysis", [
            ["Window Size (W)", "MAE (Cycles)", "Accuracy (R2)", "Std Dev (σ)", "90% Safety Bracket", "Remarks"],
            ["W = 5 Cycles", "81.68 Cycles", "81.27%", "98.90 Cycles", "±124.00 Cycles", "Too sensitive to weather changes."],
            ["W = 10 Cycles (Chosen)", "81.68 Cycles", "78.77%", "99.35 Cycles", "±122.00 Cycles", "Best historical memory depth."],
            ["W = 15 Cycles", "79.55 Cycles", "79.44%", "97.40 Cycles", "±122.00 Cycles", "Smooth calculation across normal driving."],
            ["W = 20 Cycles", "76.87 Cycles", "81.17%", "94.20 Cycles", "±119.80 Cycles", "Filters out voltage sensor noise."],
            ["W = 50 Cycles", "73.33 Cycles", "82.67%", "88.10 Cycles", "±116.40 Cycles", "Hides sudden drops in real driving."]
        ], [Inches(1.3), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.3), Inches(1.4)]),
        
        ("Table 3 LightGBM Hyperparameter Tuning Experiments", [
            ["Config", "Trees", "Leaves", "LR", "MAE Error", "Accuracy (R2)", "Std Dev (σ)", "Remarks"],
            ["Config 1 (Underfit)", "100", "15", "0.10", "90.99 Cycles", "76.65%", "111.20 Cycles", "Too simple; misses sudden drops."],
            ["Config 2 (Overfit)", "600", "127", "0.01", "75.92 Cycles", "81.23%", "92.40 Cycles", "Too heavy for cheap car chips."],
            ["Config 3 (Unregularized)", "200", "63", "0.20", "78.76 Cycles", "79.90%", "96.80 Cycles", "Jumpy countdown on rough roads."],
            ["Config 4 (Chosen)", "300", "31", "0.05", "81.35 Cycles", "78.52%", "99.35 Cycles", "Best industrial car chip choice."],
            ["Config 5 (Heavy Model)", "500", "31", "0.05", "80.17 Cycles", "78.61%", "98.10 Cycles", "Slightly better, but slower chip."]
        ], [Inches(1.4), Inches(0.6), Inches(0.6), Inches(0.5), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.3)]),
        
        ("Table 4 Checkpoint Polling Interval Evaluation Table", [
            ["Interval", "Samples", "MAE Error", "Accuracy (R2)", "Std Dev (σ)", "90% Bracket", "Remarks"],
            ["5 Cycles (Chosen)", "4,234", "81.35 Cycles", "78.52%", "99.35 Cycles", "±122.00 Cycles", "Best balance of speed and safety."],
            ["10 Cycles", "2,124", "80.68 Cycles", "79.02%", "98.12 Cycles", "±118.50 Cycles", "Good, but delays dashboard alerts."],
            ["15 Cycles", "1,421", "82.11 Cycles", "78.44%", "101.40 Cycles", "±121.70 Cycles", "Slightly less precise during drops."],
            ["20 Cycles", "1,070", "79.34 Cycles", "79.85%", "97.05 Cycles", "±116.00 Cycles", "Leaves 2-3 week blind spots."],
            ["50 Cycles (Worst)", "435", "88.54 Cycles", "77.32%", "108.90 Cycles", "±130.00 Cycles", "Poor; lags badly and misses drops."]
        ], [Inches(1.2), Inches(0.8), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.4)]),
        
        ("Table 5 Unseen Test Cohort Validation Performance", [
            ["Dataset Split", "Cell Count", "Checkpoints", "MAE Error", "R2 Accuracy"],
            ["Training Set", "100 Cells", "18,240", "48.70 cycles", "95.74%"],
            ["Unseen Test Set", "24 Cells", "4,234", "81.35 cycles", "78.52%"]
        ], [Inches(1.5), Inches(1.2), Inches(1.2), Inches(1.5), Inches(1.5)]),
        
        ("Table 6 Empirical Benchmarking of Predictive Models on Unseen Test Cohort", [
            ["Model", "MAE (Cycles)", "R2 (%)", "Flash Size", "MCU Loop", "Automotive ECU Feasibility"],
            ["Linear Regression", "152.90", "56.12%", "1.2 KB", "0.22 ms", "Deployable but high error (>150c)."],
            ["Random Forest", "80.18", "79.15%", "28.1 MB", "1.85 ms", "FAILED: Exceeds 1 MB Flash limit."],
            ["XGBoost Regressor", "82.53", "80.25%", "1.36 MB", "0.68 ms", "FAILED: Exceeds 1 MB Flash limit."],
            ["LightGBM (Ours)", "79.91", "81.62%", "809 KB", "0.41 ms", "OPTIMAL: Fits Flash, <1.5 ms loop."]
        ], [Inches(1.4), Inches(1.0), Inches(0.8), Inches(1.0), Inches(1.0), Inches(1.8)]),
        
        ("Table 7 Prognostic Maintenance Alert Classification Performance Metrics", [
            ["Performance Metric", "Mathematical Definition", "Numerical Calculation", "Result"],
            ["Accuracy", "(TP + TN) / (TP + TN + FP + FN)", "(498 + 3600) / (498 + 3600 + 72 + 64)", "96.79%"],
            ["Precision", "TP / (TP + FP)", "498 / (498 + 72)", "87.37%"],
            ["Recall", "TP / (TP + FN)", "498 / (498 + 64)", "88.61%"],
            ["F1 Score", "2 · (Precision · Recall) / (Precision + Recall)", "2 · (0.8737 · 0.8861) / (0.8737 + 0.8861)", "87.98%"]
        ], [Inches(1.3), Inches(2.2), Inches(2.2), Inches(1.0)]),
        
        ("Table 8 Hardware Microcontroller Benchmarking & Timing Breakdown", [
            ["Execution Step", "Operations Performed", "Measured Latency", "Hardware Microchip Feasibility"],
            ["1. Sensor Data Acquisition", "Reading V, I, T buffers", "0.12 ms", "Direct CAN bus register read"],
            ["2. Feature Extraction", "Rolling dQ log var (W=10)", "0.42 ms", "Simple integer rolling variance"],
            ["3. Tree Evaluation", "Traversing 300 trees", "0.41 ms", "Integer IF/ELSE threshold logic"],
            ["TOTAL INFERENCE", "Full End-to-End Prediction", "0.95 ms", "Leaves >99.9% CPU free (<1.5 ms budget)"]
        ], [Inches(1.6), Inches(1.8), Inches(1.2), Inches(2.4)])
    ]
    
    doc = docx.Document()
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Dynamic RUL Prediction of LFP Batteries\nComplete Manuscript Tables (Table 1 to Table 8)")
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(16, 44, 87) # Navy
    doc.add_paragraph() # Spacing
    
    for title, rows, widths in tables_data:
        # Heading
        h_p = doc.add_paragraph()
        h_run = h_p.add_run(title)
        h_run.font.bold = True
        h_run.font.size = Pt(12)
        h_run.font.name = "Calibri"
        h_run.font.color.rgb = RGBColor(16, 44, 87)
        
        # Table
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        for i, row_data in enumerate(rows):
            for j, val in enumerate(row_data):
                t.cell(i, j).text = val
        format_table(t, widths)
        
        # Spacing after table
        doc.add_paragraph()
        
    doc_paths = [
        r"D:\chandru downloads\LFP_RUL_Manuscript_All_Tables.docx",
        r"d:\chandru project\RUL prediction\LFP_RUL_Manuscript_All_Tables.docx"
    ]
    for dp in doc_paths:
        doc.save(dp)
        print(f"Saved Standalone Tables Document to: {dp}")
    print("\nALL TASKS COMPLETED SUCCESSFULLY!")

if __name__ == "__main__":
    create_figures_zip()
    create_all_tables_doc()
