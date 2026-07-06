import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX

def set_font(p, font_name="Times New Roman", size_pt=12, bold=False, italic=False, color=RGBColor(0,0,0), highlight=None):
    for run in p.runs:
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
        if highlight is not None:
            run.font.highlight_color = highlight

# Section 5 text with precise IEEE citations added without changing a single word!
SEC_TITLE = "5. Methodology and Validation Framework"

SUB_5_1 = "5.1 LightGBM Model Architecture and Tabular Rationale"
PARA_5_1_A = "Prognostic models for automotive embedded systems need to be effective in terms of both predictive performance and efficiency [10], [18]. Despite the success of Long Short-Term Memory networks in estimating the lifetime of batteries [7], [8], they employ complex matrix multiplication for computing input streams [8]. Deep recurrent networks require excessive memory overhead and inference times in automotive control units [8], [10]. In contrast, gradient boosting techniques of decision trees, especially in the case of LightGBM model, have shown better performance when applied to structured tabular data [10], [11], [18]. LightGBM operates on histogram-based binning of continuous variables and performs tree building in a leaf-wise fashion with depth limit, which leads to a much faster inference than neural networks [11], [18]."
PARA_5_1_B = "To improve the efficiency of the model in the presence of noise in automotive hardware, a set of hyperparameter tuning tests was conducted using five different configurations [10], [18]. As can be seen in Table 3, the first configuration with 100 trees, 15 leaves, and learning rate 0.10 was a typical case of serious underfitting because its MAE 90.99 cycles failed to represent sharp drops in capacity [1]. However, Config 2 with 600 trees and 127 leaves suffered from overfitting of noise in training data and created complex models that were not suitable for automotive hardware [10], [18]. Finally, Config 4 with 300 estimators, 31 maximum leaves, and a low learning rate of 0.05 was selected as the best industrial configuration with 78.52% of accuracy (MAE 81.35) [11], [18]."
TABLE_3_TEXT = "Table 3 LightGBM Hyperparameter Tuning Experiments"

SUB_5_2 = "5.2 Grouped Leave-Cells-Out Validation"
PARA_5_2 = "There is a considerable leakage in time-series battery data using the classical k-fold cross-validation procedure [5]. When the same physical battery’s cycles are assigned to training and testing sets randomly, the model acquires high accuracy through interpolation between adjacent cycles [5]. In order to prevent any leakage, GroupShuffleSplit validation is employed at the level of physical cells [1], [5]. In particular, 124 commercial LFP cells are divided into 100 cells assigned to training and 24 cells assigned to testing [1]. As can be seen from the training flow architecture (Fig. 2), the testing vault is out of bounds during model training and tuning [1], [5]."
FIG_2_TEXT = "Fig. 2 Training and validation flow architecture illustrating leakage-free GroupShuffleSplit at the physical cell level. The 24 test evaluation batteries remain strictly unseen during LightGBM model training, serving as an unbiased out-of-sample benchmark."

SUB_5_3 = "5.3 Conformal Prediction for Uncertainty Quantification"
PARA_5_3 = "The point estimates of remaining useful life do not have sufficient predictive value for safety-critical algorithms for vehicle controllers without mathematics-based uncertainty bounds [14], [15]. For reliable systems, the technique makes use of split conformal prediction to create valid prediction intervals [14], [19]. In split conformal prediction, the nonconformity scores are tuned using the out-of-fold calibration errors that are independent of the 24 test cells [14], [19]. Specifically, the absolute prediction errors in the validation cell are used to compute the empirical error quantile of 90th percentile as 122 cycles [19]. The calibrated bound creates a safety interval of ±122 cycles around the LightGBM point estimate [19]. Unlike the Bayesian neural networks with the assumption of Gaussian error distribution [15], conformal prediction ensures valid coverage even with trajectory skewness [14], [19]."

SUB_5_4 = "5.4 Checkpoint Polling Interval Selection"
PARA_5_4_A = "The problem of how often the diagnostics of polling is performed by the Battery Management System onboard relates to the tradeoff between computation of the microcontroller and trajectory reactivity [10], [18]. Diagnostics in each cycle is the unnecessary computation waste on static plateaus, but too few polls will cause missing diagnostics [10], [18]. Below is presented the analysis of polling intervals done for 22,474 checkpoints with 5, 10, 15, 20, and 50 cycles [1], [18]."
TABLE_4_TEXT = "Table 4 Checkpoint Polling Interval Evaluation Table"
PARA_5_4_B = "Table 4 and Figure 3 illustrate convincingly that the 5-cycle polling resolution is an optimal solution balancing performance and responsiveness [10], [18]. Although the very scarce 20-cycle polling approach provides some advantage regarding the mean error in the situation of stable plateaus, it produces dangerous “blind zones” during one to three weeks of road operation under realistic conditions [10], [18]. Toward the end of the life cycle, LFP batteries experience non-linear capacity losses when their capacity drops to the threshold level of 80% within 15 cycles [1], [16]. As a result, the 20-cycle polling resolution enables the battery cell to die without being detected by the algorithm [10], [18]."
FIG_3_TEXT = "Fig. 3 Comparison of diagnostic responsiveness between high-resolution 5-cycle polling (left) and sparse 20-cycle polling (right). Polling every 20 cycles introduces a severe inspection blind spot during the non-linear aging plunge, allowing a cell to drop below the safety threshold unmonitored"

def add_section5_content(doc_or_para, is_insert_before=False):
    def add_p(text, space_before=0, space_after=6, bold=False, italic=False, size_pt=12, align=WD_ALIGN_PARAGRAPH.LEFT, highlight=None):
        if is_insert_before:
            p = doc_or_para.insert_paragraph_before()
        else:
            p = doc_or_para.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)
        if highlight is not None:
            run.font.highlight_color = highlight
        return p

    # Section 5 Title
    add_p(SEC_TITLE, space_before=18, space_after=6, bold=True, size_pt=14)
    
    # 5.1
    add_p(SUB_5_1, space_before=12, space_after=4, bold=True, size_pt=12)
    add_p(PARA_5_1_A, space_before=0, space_after=6, size_pt=12)
    add_p(PARA_5_1_B, space_before=0, space_after=6, size_pt=12)
    add_p(TABLE_3_TEXT, space_before=6, space_after=12, italic=True, size_pt=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, highlight=WD_COLOR_INDEX.YELLOW)
    
    # 5.2
    add_p(SUB_5_2, space_before=12, space_after=4, bold=True, size_pt=12)
    add_p(PARA_5_2, space_before=0, space_after=6, size_pt=12)
    add_p(FIG_2_TEXT, space_before=6, space_after=12, italic=True, size_pt=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, highlight=WD_COLOR_INDEX.YELLOW)
    
    # 5.3
    add_p(SUB_5_3, space_before=12, space_after=4, bold=True, size_pt=12)
    add_p(PARA_5_3, space_before=0, space_after=6, size_pt=12)
    
    # 5.4
    add_p(SUB_5_4, space_before=12, space_after=4, bold=True, size_pt=12)
    add_p(PARA_5_4_A, space_before=0, space_after=6, size_pt=12)
    add_p(TABLE_4_TEXT, space_before=6, space_after=12, italic=True, size_pt=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, highlight=WD_COLOR_INDEX.YELLOW)
    add_p(PARA_5_4_B, space_before=0, space_after=6, size_pt=12)
    add_p(FIG_3_TEXT, space_before=6, space_after=12, italic=True, size_pt=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, highlight=WD_COLOR_INDEX.YELLOW)

def create_standalone_section5(out_path):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    add_section5_content(doc, is_insert_before=False)
    doc.save(out_path)
    print(f"Saved standalone Section 5 to: {out_path}")

def update_master_manuscript(out_master_path):
    # Try reading from Final_Manuscript_with_Section4.docx first
    source_path = r"D:\chandru project\Final_Manuscript_with_Section4.docx"
    if not os.path.exists(source_path):
        source_path = r"D:\chandru project\Final_Manuscript_with_Section3.docx"
    if not os.path.exists(source_path):
        source_path = r"D:\chandru project\Final_Manuscript_Abstract_Intro_RelatedWork.docx"
    
    if not os.path.exists(source_path):
        print(f"Source file not found: {source_path}")
        return
        
    doc = docx.Document(source_path)
    print(f"Opened master source: {source_path} (Total paras: {len(doc.paragraphs)})")
    
    # Find where References section starts
    ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "References" or p.text.strip().startswith("REFERENCES"):
            ref_idx = i
            break
            
    if ref_idx is None:
        print("Could not find References section, appending Section 5 to end.")
        add_section5_content(doc, is_insert_before=False)
    else:
        print(f"Found References section at paragraph index {ref_idx}. Inserting Section 5 before References.")
        ref_p = doc.paragraphs[ref_idx]
        add_section5_content(ref_p, is_insert_before=True)
        
    # Save to out_master_path
    doc.save(out_master_path)
    print(f"Updated master document with Section 5 saved to: {out_master_path}")
    
    # Also try saving back to source_path if not locked
    try:
        doc.save(source_path)
        print(f"Updated original source document: {source_path}")
    except PermissionError:
        print(f"Note: {source_path} is open in Word, so saved to {out_master_path}")

    # Save backup in reports
    backup_path = r"d:\chandru project\RUL prediction\reports\Final_Manuscript_with_Section5.docx"
    doc.save(backup_path)
    print(f"Saved backup to: {backup_path}")

if __name__ == "__main__":
    standalone_path = r"D:\chandru project\Section5_Methodology_and_Validation_Framework.docx"
    create_standalone_section5(standalone_path)
    
    update_master_manuscript(r"D:\chandru project\Final_Manuscript_with_Section5.docx")
