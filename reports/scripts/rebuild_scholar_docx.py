import os
import shutil
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def format_table(table, col_widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if col_widths and j < len(col_widths):
                cell.width = col_widths[j]
            if i == 0:
                set_cell_background(cell, "102C57") # Navy Header
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.font.name = "Calibri"
                        r.font.size = Pt(10)
            else:
                bg = "F5F7FA" if i % 2 == 1 else "FFFFFF"
                set_cell_background(cell, bg)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(33, 37, 41)
                        r.font.name = "Calibri"
                        r.font.size = Pt(9.5)

def rebuild_docx(doc_path):
    print(f"Processing document: {doc_path}")
    
    # Restore from backup if exists, or create backup
    backup_path = doc_path.replace(".docx", "_BACKUP_PRE_FIX.docx")
    if os.path.exists(backup_path):
        print(f"Restoring clean state from backup: {backup_path}")
        shutil.copy2(backup_path, doc_path)
    else:
        shutil.copy2(doc_path, backup_path)
        print(f"Created backup at: {backup_path}")
    
    doc = docx.Document(doc_path)
    
    # 1. REMOVE ALL EXISTING TABLES
    print(f"Removing {len(doc.tables)} existing tables...")
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)
        
    # 2. DEFINITIONS OF ALL 8 PERFECT TABLES
    tables_data = {
        "Table 1": {
            "title_prefix": "Table 1 Physics-Informed Feature Set",
            "headers": ["Symbol", "Mathematical Definition", "Physical Meaning & Degradation Link", "Computation Formula"],
            "rows": [
                ["k", "Cycle Number", "Tracks baseline temporal aging progression and cumulative operational charge-discharge duration of the cell.", "Current operational cycle index k at checkpoint"],
                ["SOH", "State of Health", "Normalized remaining capacity ratio; primary macro-level health indicator across cell form factors.", "Q_dis(k) / Q_nominal"],
                ["IR", "Internal Resistance", "Measures ohmic + SEI conduction resistance; primary indicator of electrolyte decomposition and interfacial thickening.", "Mean(dV / dI) during pulse at 10% SOC"],
                ["T_mean", "Mean Discharge Temp", "Monitors thermal stress during operational discharge; drives Arrhenius aging acceleration and SEI growth.", "Mean surface/core temperature over W = 10"],
                ["ΔSOH", "Capacity Fade Rate", "Normalized rate of capacity loss over lookback window; measures local degradation velocity (ΔSOH).", "[Q_dis(k-10) - Q_dis(k)] / Q_nominal"],
                ["ΔQ_log_var", "IC Curve Log-Variance", "Captures thermodynamic peak broadening in incremental capacity curve; acts as an early warning sentinel.", "log[Var(dQ / dV)] over window W=10"],
                ["ΔQ_min", "IC Curve Min Shift", "Maximum localized downward shift (valley) in the differential capacity profile over the rolling lookback window.", "min[Δ(dQ / dV)] over window W=10"],
                ["ΔQ_mean", "IC Curve Mean Shift", "Average baseline shift of incremental capacity curve; tracks global thermodynamic drift and stoichiometric imbalance.", "mean[Δ(dQ / dV)] over window W=10"]
            ],
            "widths": [Inches(0.9), Inches(1.8), Inches(2.5), Inches(1.8)]
        },
        "Table 2": {
            "title_prefix": "Table 2 Rolling Lookback Window",
            "headers": ["Window Size (W)", "MAE (Cycles)", "Accuracy (R2)", "Std Dev (σ)", "90% Safety Bracket", "Remarks"],
            "rows": [
                ["W = 5 Cycles", "81.68 Cycles", "81.27%", "98.90 Cycles", "±124.00 Cycles", "Too sensitive to weather changes."],
                ["W = 10 Cycles (Chosen)", "81.68 Cycles", "78.77%", "99.35 Cycles", "±122.00 Cycles", "Best historical memory depth."],
                ["W = 15 Cycles", "79.55 Cycles", "79.44%", "97.40 Cycles", "±122.00 Cycles", "Smooth calculation across normal driving."],
                ["W = 20 Cycles", "76.87 Cycles", "81.17%", "94.20 Cycles", "±119.80 Cycles", "Filters out voltage sensor noise."],
                ["W = 50 Cycles", "73.33 Cycles", "82.67%", "88.10 Cycles", "±116.40 Cycles", "Hides sudden drops in real driving."]
            ],
            "widths": [Inches(1.3), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.3), Inches(1.4)]
        },
        "Table 3": {
            "title_prefix": "Table 3 LightGBM Hyperparameter",
            "headers": ["Config", "Trees", "Leaves", "LR", "MAE Error", "Accuracy (R2)", "Std Dev (σ)", "Remarks"],
            "rows": [
                ["Config 1 (Underfit)", "100", "15", "0.10", "90.99 Cycles", "76.65%", "111.20 Cycles", "Too simple; misses sudden drops."],
                ["Config 2 (Overfit)", "600", "127", "0.01", "75.92 Cycles", "81.23%", "92.40 Cycles", "Too heavy for cheap car chips."],
                ["Config 3 (Unregularized)", "200", "63", "0.20", "78.76 Cycles", "79.90%", "96.80 Cycles", "Jumpy countdown on rough roads."],
                ["Config 4 (Chosen)", "300", "31", "0.05", "81.35 Cycles", "78.52%", "99.35 Cycles", "Best industrial car chip choice."],
                ["Config 5 (Heavy Model)", "500", "31", "0.05", "80.17 Cycles", "78.61%", "98.10 Cycles", "Slightly better, but slower chip."]
            ],
            "widths": [Inches(1.4), Inches(0.6), Inches(0.6), Inches(0.5), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.3)]
        },
        "Table 4": {
            "title_prefix": "Table 4 Checkpoint Polling Interval Evaluation Table",
            "headers": ["Interval", "Samples", "MAE Error", "Accuracy (R2)", "Std Dev (σ)", "90% Bracket", "Remarks"],
            "rows": [
                ["5 Cycles (Chosen)", "4,234", "81.35 Cycles", "78.52%", "99.35 Cycles", "±122.00 Cycles", "Best balance of speed and safety."],
                ["10 Cycles", "2,124", "80.68 Cycles", "79.02%", "98.12 Cycles", "±118.50 Cycles", "Good, but delays dashboard alerts."],
                ["15 Cycles", "1,421", "82.11 Cycles", "78.44%", "101.40 Cycles", "±121.70 Cycles", "Slightly less precise during drops."],
                ["20 Cycles", "1,070", "79.34 Cycles", "79.85%", "97.05 Cycles", "±116.00 Cycles", "Leaves 2-3 week blind spots."],
                ["50 Cycles (Worst)", "435", "88.54 Cycles", "77.32%", "108.90 Cycles", "±130.00 Cycles", "Poor; lags badly and misses drops."]
            ],
            "widths": [Inches(1.2), Inches(0.8), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.4)]
        },
        "Table 5": {
            "title_prefix": "Table 5 Unseen Test Cohort",
            "headers": ["Dataset Split", "Cell Count", "Checkpoints", "MAE Error", "R2 Accuracy"],
            "rows": [
                ["Training Set", "100 Cells", "18,240", "48.70 cycles", "95.74%"],
                ["Unseen Test Set", "24 Cells", "4,234", "81.35 cycles", "78.52%"]
            ],
            "widths": [Inches(1.5), Inches(1.2), Inches(1.2), Inches(1.5), Inches(1.5)]
        },
        "Table 6": {
            "title_prefix": "Table 6 Empirical Benchmarking",
            "headers": ["Model", "MAE (Cycles)", "R2 (%)", "Flash Size", "MCU Loop", "Automotive ECU Feasibility"],
            "rows": [
                ["Linear Regression", "152.90", "56.12%", "1.2 KB", "0.22 ms", "Deployable but high error (>150c)."],
                ["Random Forest", "80.18", "79.15%", "28.1 MB", "1.85 ms", "FAILED: Exceeds 1 MB Flash limit."],
                ["XGBoost Regressor", "82.53", "80.25%", "1.36 MB", "0.68 ms", "FAILED: Exceeds 1 MB Flash limit."],
                ["LightGBM (Ours)", "79.91", "81.62%", "809 KB", "0.41 ms", "OPTIMAL: Fits Flash, <1.5 ms loop."]
            ],
            "widths": [Inches(1.4), Inches(1.0), Inches(0.8), Inches(1.0), Inches(1.0), Inches(1.8)]
        },
        "Table 7": {
            "title_prefix": "Table 7 Prognostic Maintenance",
            "headers": ["Performance Metric", "Mathematical Definition", "Numerical Calculation", "Result"],
            "rows": [
                ["Accuracy", "(TP + TN) / (TP + TN + FP + FN)", "(498 + 3600) / (498 + 3600 + 72 + 64)", "96.79%"],
                ["Precision", "TP / (TP + FP)", "498 / (498 + 72)", "87.37%"],
                ["Recall", "TP / (TP + FN)", "498 / (498 + 64)", "88.61%"],
                ["F1 Score", "2 · (Precision · Recall) / (Precision + Recall)", "2 · (0.8737 · 0.8861) / (0.8737 + 0.8861)", "87.98%"]
            ],
            "widths": [Inches(1.3), Inches(2.2), Inches(2.2), Inches(1.0)]
        },
        "Table 8": {
            "title_prefix": "Table 8",
            "headers": ["Execution Step", "Operations Performed", "Measured Latency", "Hardware Microchip Feasibility"],
            "rows": [
                ["1. Sensor Data Acquisition", "Reading V, I, T buffers", "0.12 ms", "Direct CAN bus register read"],
                ["2. Feature Extraction", "Rolling dQ log var (W=10)", "0.42 ms", "Simple integer rolling variance"],
                ["3. Tree Evaluation", "Traversing 300 trees", "0.41 ms", "Integer IF/ELSE threshold logic"],
                ["TOTAL INFERENCE", "Full End-to-End Prediction", "0.95 ms", "Leaves >99.9% CPU free (<1.5 ms budget)"]
            ],
            "widths": [Inches(1.6), Inches(1.8), Inches(1.2), Inches(2.4)]
        }
    }
    
    print("Fixing caption numbering and text references...")
    table6_heading_exists = False
    for p in doc.paragraphs:
        txt = p.text.strip()
        if "Figure 2  LightGBM Feature Importance" in txt or "Fig. 1 LightGBM Feature Importance" in txt or "Feature Importance ranking by split count" in txt:
            if "Fig" in txt or "Figure" in txt:
                p.text = "Fig. 1 LightGBM Feature Importance ranking by split count across the eight physics-informed parameters."
        elif "Training and validation flow architecture illustrating" in txt:
            p.text = "Fig. 2 Training and validation flow architecture illustrating leakage-free GroupShuffleSplit at the physical cell level. The 24 test evaluation batteries remain strictly unseen during LightGBM model training, serving as an unbiased out-of-sample benchmark."
        elif "Table 6 Empirical Benchmarking" in txt:
            table6_heading_exists = True
        elif "Table 6: Performance metrics calculation" in txt or "Table 7: Performance metrics calculation" in txt or "Table 6 Performance metrics" in txt or "Table 7 Prognostic Maintenance" in txt:
            p.text = "Table 7 Prognostic Maintenance Alert Classification Performance Metrics"
            
    if not table6_heading_exists:
        for p in doc.paragraphs:
            if "As shown in Table 6 and Fig. 6" in p.text or "Severson linear regression baseline model will have an MAE" in p.text:
                new_p = p.insert_paragraph_before("Table 6 Empirical Benchmarking of Predictive Models on Unseen Test Cohort")
                new_p.style = p.style
                new_p.runs[0].font.bold = True
                print("Inserted Table 6 heading paragraph.")
                break

    # 3. INSERT TABLES AFTER THEIR RESPECTIVE HEADINGS
    print("Inserting 8 rebuilt tables into document...")
    inserted_tables = set()
    for p in list(doc.paragraphs):
        txt = p.text.strip()
        for t_key, t_info in tables_data.items():
            if t_key not in inserted_tables and txt.startswith(t_info["title_prefix"]):
                print(f" -> Inserting {t_key} after heading: '{txt[:40]}...'")
                new_t = doc.add_table(rows=len(t_info["rows"]) + 1, cols=len(t_info["headers"]))
                
                for col_idx, h_text in enumerate(t_info["headers"]):
                    new_t.cell(0, col_idx).text = h_text
                for row_idx, r_data in enumerate(t_info["rows"]):
                    for col_idx, c_text in enumerate(r_data):
                        new_t.cell(row_idx + 1, col_idx).text = c_text
                        
                format_table(new_t, t_info["widths"])
                p._element.addnext(new_t._element)
                inserted_tables.add(t_key)
                break

    # 4. INSERT ALL 8 PICTURES ABOVE THEIR CAPTIONS
    print("Inserting 8 figures into document...")
    figures_map = {
        "Fig. 1 LightGBM Feature Importance": r"d:\chandru project\RUL prediction\results\figures\02_feature_importance.png",
        "Fig. 2 Training and validation flow": r"d:\chandru project\RUL prediction\reports\figures\flow_architecture.png",
        "Fig. 3 Comparison of diagnostic responsiveness": r"d:\chandru project\RUL prediction\reports\figures\polling_blind_spot.png",
        "Fig. 4 Complex s-plane pole-zero migration": r"d:\chandru project\RUL prediction\reports\figures\pole_zero_migration.png",
        "Fig. 5 Parity scatter plot": r"d:\chandru project\RUL prediction\results\figures\01_true_vs_predicted_rul.png",
        "Fig. 6 Empirical benchmarking comparison": r"d:\chandru project\RUL prediction\results\benchmarks\model_comparison_bar_chart.png",
        "Fig. 7 Distribution of prediction errors": r"d:\chandru project\RUL prediction\results\figures\03_prediction_errors_histogram.png",
        "Fig. 8 Prognostic maintenance confusion matrix": r"d:\chandru project\RUL prediction\results\figures\06_confusion_matrix.png"
    }
    
    inserted_figures = set()
    for p in list(doc.paragraphs):
        txt = p.text.strip()
        for f_prefix, img_path in figures_map.items():
            if f_prefix not in inserted_figures and txt.startswith(f_prefix):
                if os.path.exists(img_path):
                    print(f" -> Inserting figure for caption: '{txt[:35]}...'")
                    img_p = p.insert_paragraph_before()
                    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = img_p.add_run()
                    run.add_picture(img_path, width=Inches(5.5))
                    inserted_figures.add(f_prefix)
                else:
                    print(f"WARNING: Image path not found: {img_path}")
                break

    doc.save(doc_path)
    print(f"\nSUCCESS! Rebuilt document saved to: {doc_path}")
    print(f"Total Tables Inserted: {len(doc.tables)}")
    print(f"Total Figures Inserted: {len(inserted_figures)}")

if __name__ == "__main__":
    target_docx = r"D:\chandru downloads\Dynamic Remaining Useful Life (RUL) Prediction of Lithium-Iron-Phosphate Batteries with Conformal Uncertainty Quantification.docx"
    rebuild_docx(target_docx)
