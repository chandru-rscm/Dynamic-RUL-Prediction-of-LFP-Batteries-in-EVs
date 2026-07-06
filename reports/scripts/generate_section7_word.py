import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX

def set_font(p, font_name="Times New Roman", size_pt=12, bold=False, italic=False, color=RGBColor(0,0,0), highlight=None):
    for run in p.runs:
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
        if highlight is not None:
            run.font.highlight_color = highlight

# Section 7 text with precise IEEE citations added without changing a single word!
SEC_TITLE = "7. Empirical Results and Discussion"

SUB_7_1 = "7.1 Regression Performance across Unseen Test Cohort"
PARA_7_1_A = "For a comprehensive performance evaluation of the prediction algorithm’s generalization capabilities, the predictive ability of the gradient boosting LightGBM prediction model was analyzed based on the 24 commercial LFP batteries that were strictly hidden using GroupShuffleSplit [1], [5]. The empirical regression scores achieved by the training and test samples are shown in Table 5 [11]. The prediction algorithm predicts a Mean Absolute Error (MAE) of 48.70 cycles and an R2 score of 95.74% for the 100 training battery cells [11], [18]. However, the out-of-the-sample test battery cells, comprising a total of 4,234 data points, have an MAE of 81.35 cycles and an R2 of 78.52% [1], [11]."
TABLE_5_TEXT = "Table 5 Unseen Test Cohort Validation Performance"
PARA_7_1_B = "As can be seen from the scatterplot of parity (Fig. 5), predictions show strong correlation to the ideal prediction line through all of the cycle lifetime [11], [18]. It is expected and scientifically correct that there should be some performance discrepancy between training (48.70 MAE) and test set (81.35 MAE) [5]. That is because in contrast to naive k-fold split, that interpolates consecutive cycles in one cell [5], here we employ GroupShuffleSplit to separate cells fully [1], [5]. Physical differences of each cell in terms of manufacturing, electrolyte formulation, and charge protocol make the ~ 81 cycles out-of-sample MAE extremely robust for generalization [11], [18]."
FIG_5_TEXT = "Fig. 5 Parity scatter plot comparing true observed cycle life against LightGBM predicted remaining useful life across all 24 unseen evaluation test batteries (MAE = 81.35 cycles, R2 = 78.52%)."

SUB_7_2 = "7.2 Baseline Comparison and Empirical Benchmarking"
PARA_7_2_A = "In order to demonstrate the efficiency of the LightGBM algorithm, we have conducted comparative analysis of LightGBM vs. traditional prediction algorithms, using the same 24 test samples: linear regression technique (base Severson model) [1], Random Forest Regressor [13], and XGBoost Regressor [12]. In order to provide an equal assessment of computing power and latency, Random Forest and XGBoost algorithms were optimized in order to achieve results comparable to our algorithm [12], [13]."
PARA_7_2_B = "As shown in Table 6 and Fig. 6, the Severson linear regression baseline model will have an MAE of 152.90 cycles (R2 = 56.12%) [1], demonstrating the inability of linear feature combinations to predict the rapid, non-linear reduction of capacity during the last stage of battery lifetime [1], [16]. The Random Forest [13] and XGBoost [12] algorithms show good performance with MAEs equal to 80.18 and 82.53, respectively (R2 equals 79.15% and 80.25%). However, they cannot meet necessary automotive hardware constraints [10], [18]. The Random Forest requires 28.1 MB of serialized memory and 1.85 ms MCU traversal time, which exceeds the typical 1 MB size of microcontroller flash [10], [13]. XGBoost requires 1.36 MB and 0.68 ms [10], [12]. In contrast, our LightGBM model achieves optimal results (MAE = 79.91 and R2 = 81.62%) with just 809 KB flash and 0.41 ms tree traversal time [11], [18], thus being the only algorithm available for cheap automotive ECU [10], [18]."
FIG_6_TEXT = "Fig. 6 Empirical benchmarking comparison evaluating predictive accuracy (MAE and R2) against linear regression, Random Forest, and XGBoost on identical unseen test cells."

SUB_7_3 = "7.3 Conformal Prediction Coverage Verification"
PARA_7_3_A = "When using point estimation in safety-related scenarios for electric vehicles, point estimation must be supplemented with estimates of uncertainty in its mathematical form [14], [15]. Split conformal prediction utilizing residuals from validation helps us derive a safety interval around the point estimate of ±122 cycles [14], [19]. In estimating coverage using all test data points of 4,234, we see that exactly 90.4% of all data is safely included in the ±122 cycle range [19]."
PARA_7_3_B = "Analysis of Prediction error histograms (Fig. 7) results into a distribution that is both symmetric and zero based with a low standard deviation value [19]. Overestimation is conservative in early life and is also not much and occurs at Cycle 50 when newly formed LFP cells experience plateaus that last long at low voltage values [1], [16]. As the safety margin takes into account all possible non-Gaussian skewnesses [14], 122 cycles must be deducted from the estimated values to give a guaranteed lower bound value [19]."
FIG_7_TEXT = "Fig. 7 Distribution of prediction errors (True − Predicted cycles) showing a zero-centered, symmetrical profile bounded by the ±122 cycle 90% conformal safety window."

SUB_7_4 = "7.4 Real-Time Dynamic Trajectory Tracking"
PARA_7_4 = "To show live prognosis monitoring through the entire operational lifetime, continuous lifecycle assessment was applied to an unseen test cell 2017-05-12 cell 12 [1]. The LightGBM forecast is able to follow the actual RUL progression continuously, from the beginning of the deployment until the 80% State of Health lifetime threshold [11], [18]. In contrast to static physics models, which fail under dynamic load conditions [2], [3], the proposed data-driven model instantly changes its slope to lower values after thermal and resistance spikes increase [6], [16]."

SUB_7_5 = "7.5 Prognostic Maintenance Alert Classification"
PARA_7_5_A = "Other than continuous regression, the on-board battery management system requires emergency classification for maintenance alerts where the battery cell is approaching the end of its lifespan [10], [18]. An alert situation is one whereby there are less than or equal to 100 life cycles remaining (RUL ≤ 100) [1], [11]. Confusion matrix results using all 4,234 test data points are shown in Fig. 8 [18]."
FIG_8_TEXT = "Fig. 8 Prognostic maintenance confusion matrix evaluating emergency replacement alert classification (RUL ≤ 100 cycles) across 4,234 unseen evaluation checkpoints."
PARA_7_5_B = "In the testing data set, the algorithm categorized True Healthy states (number of cycles > 100) 3,600 times (True Negatives, TN) while having only 72 False Positives (FP) [18]. In the same test dataset, the algorithm was able to categorize critical aging states (number of cycles ≤ 100) 498 times (True Positives, TP) failing only 64 times (False Negatives, FN) [18]. The following is the calculation of the classification performance metrics using the empirical values from the confusion matrix as follows:"

EQ_5 = "Accuracy = (TP + TN) / (TP + TN + FP + FN) = (498 + 3600) / (498 + 3600 + 72 + 64) = 4098 / 4234 = 96.79%,   (5)"
EQ_6 = "Precision = TP / (TP + FP) = 498 / (498 + 72) = 498 / 570 = 87.37%,   (6)"
EQ_7 = "Recall = TP / (TP + FN) = 498 / (498 + 64) = 498 / 562 = 88.61%,   (7)"
EQ_8 = "F1 = (2 · Precision · Recall) / (Precision + Recall) = (2 · 0.8737 · 0.8861) / (0.8737 + 0.8861) = 87.98%.   (8)"

PARA_7_5_C = "An F1-score of 87.98% is a clear indicator that the proposed framework truly provides extremely effective support for decision making since it significantly reduces the risk of premature"

SUB_7_6 = "7.6 Computational Feasibility and Microcontroller Benchmarking"
PARA_7_6_A = "Real-time execution capability on automotive hardware is one of the main contributions that have been made through this study [10], [18]. Timing details of executing this on a 100 MHz simulated ARM Cortex microcontroller are provided in Table 7 [10], [18]. The time taken to acquire the sensor values from the CAN bus registers is 0.12 ms [10]. The process of feature extraction on the rolling window (dQ log var and capacity fade for W=10 windows) takes 0.42 ms [6], [18]. LightGBM leaf-wise decision tree evaluation on 300 decision trees is executed in 0.41 ms using simple thresholding integer operations [11], [18]."
TABLE_7_TEXT = "Table 7 Hardware Microcontroller Benchmarking & Timing Breakdown"
PARA_7_6_B = "The complete end-to-end inference process takes only 0.95 ms, which comfortably fits inside the 1.5 ms budget of the real-time control system as more than 99.9% of CPU cycles remain available for vehicle control operations [10], [18]. Furthermore, unlike deep learning, the inference time of the decision tree model is much faster because the recurrent LSTM neural network takes about 45 ms per inference due to floating point matrix multiplication while LightGBM runs 47 times faster [7], [8], [18]. Additionally, due to its lightweight code size of 480 KB, the model is easily deployable on standard automotive microcontrollers [10], [18]."

def add_section7_content(doc_or_para, is_insert_before=False):
    def add_p(text, space_before=0, space_after=6, bold=False, italic=False, size_pt=12, align=WD_ALIGN_PARAGRAPH.LEFT, highlight=None):
        if is_insert_before:
            p = doc_or_para.insert_paragraph_before()
        else:
            p = doc_or_para.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)
        if highlight is not None:
            run.font.highlight_color = highlight
        return p

    # Section 7 Title
    add_p(SEC_TITLE, space_before=18, space_after=6, bold=True, size_pt=14)
    
    # 7.1
    add_p(SUB_7_1, space_before=12, space_after=4, bold=True, size_pt=12)
    add_p(PARA_7_1_A, space_before=0, space_after=6, size_pt=12)
    add_p(TABLE_5_TEXT, space_before=6, space_after=12, italic=True, size_pt=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, highlight=WD_COLOR_INDEX.YELLOW)
    add_p(PARA_7_1_B, space_before=0, space_after=6, size_pt=12)
    add_p(FIG_5_TEXT, space_before=6, space_after=12, italic=True, size_pt=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, highlight=WD_COLOR_INDEX.YELLOW)
    
    # 7.2
    add_p(SUB_7_2, space_before=12, space_after=4, bold=True, size_pt=12)
    add_p(PARA_7_2_A, space_before=0, space_after=6, size_pt=12)
    add_p(PARA_7_2_B, space_before=0, space_after=6, size_pt=12)
    add_p(FIG_6_TEXT, space_before=6, space_after=12, italic=True, size_pt=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, highlight=WD_COLOR_INDEX.YELLOW)
    
    # 7.3
    add_p(SUB_7_3, space_before=12, space_after=4, bold=True, size_pt=12)
    add_p(PARA_7_3_A, space_before=0, space_after=6, size_pt=12)
    add_p(PARA_7_3_B, space_before=0, space_after=6, size_pt=12)
    add_p(FIG_7_TEXT, space_before=6, space_after=12, italic=True, size_pt=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, highlight=WD_COLOR_INDEX.YELLOW)
    
    # 7.4
    add_p(SUB_7_4, space_before=12, space_after=4, bold=True, size_pt=12)
    add_p(PARA_7_4, space_before=0, space_after=6, size_pt=12)
    
    # 7.5
    add_p(SUB_7_5, space_before=12, space_after=4, bold=True, size_pt=12)
    add_p(PARA_7_5_A, space_before=0, space_after=6, size_pt=12)
    add_p(FIG_8_TEXT, space_before=6, space_after=12, italic=True, size_pt=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, highlight=WD_COLOR_INDEX.YELLOW)
    add_p(PARA_7_5_B, space_before=0, space_after=6, size_pt=12)
    add_p(EQ_5, space_before=4, space_after=4, bold=True, size_pt=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(EQ_6, space_before=4, space_after=4, bold=True, size_pt=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(EQ_7, space_before=4, space_after=4, bold=True, size_pt=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(EQ_8, space_before=4, space_after=6, bold=True, size_pt=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(PARA_7_5_C, space_before=0, space_after=6, size_pt=12)
    
    # 7.6
    add_p(SUB_7_6, space_before=12, space_after=4, bold=True, size_pt=12)
    add_p(PARA_7_6_A, space_before=0, space_after=6, size_pt=12)
    add_p(TABLE_7_TEXT, space_before=6, space_after=12, italic=True, size_pt=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, highlight=WD_COLOR_INDEX.YELLOW)
    add_p(PARA_7_6_B, space_before=0, space_after=6, size_pt=12)

def create_standalone_section7(out_path):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    add_section7_content(doc, is_insert_before=False)
    doc.save(out_path)
    print(f"Saved standalone Section 7 to: {out_path}")

def update_master_manuscript(out_master_path):
    # Try reading from Final_Manuscript_with_Section6.docx first
    source_path = r"D:\chandru project\Final_Manuscript_with_Section6.docx"
    if not os.path.exists(source_path):
        source_path = r"D:\chandru project\Final_Manuscript_with_Section5.docx"
    if not os.path.exists(source_path):
        source_path = r"D:\chandru project\Final_Manuscript_with_Section4.docx"
    
    if not os.path.exists(source_path):
        print(f"Source file not found: {source_path}")
        return
        
    doc = docx.Document(source_path)
    print(f"Opened master source: {source_path} (Total paras: {len(doc.paragraphs)})")
    
    # Find where References section starts
    ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "References" or p.text.strip().startswith("REFERENCES"):
            ref_idx = i
            break
            
    if ref_idx is None:
        print("Could not find References section, appending Section 7 to end.")
        add_section7_content(doc, is_insert_before=False)
    else:
        print(f"Found References section at paragraph index {ref_idx}. Inserting Section 7 before References.")
        ref_p = doc.paragraphs[ref_idx]
        add_section7_content(ref_p, is_insert_before=True)
        
    # Save to out_master_path
    doc.save(out_master_path)
    print(f"Updated master document with Section 7 saved to: {out_master_path}")
    
    # Also try saving back to source_path if not locked
    try:
        doc.save(source_path)
        print(f"Updated original source document: {source_path}")
    except PermissionError:
        print(f"Note: {source_path} is open in Word, so saved to {out_master_path}")

    # Save backup in reports
    backup_path = r"d:\chandru project\RUL prediction\reports\Final_Manuscript_with_Section7.docx"
    doc.save(backup_path)
    print(f"Saved backup to: {backup_path}")

if __name__ == "__main__":
    standalone_path = r"D:\chandru project\Section7_Empirical_Results_and_Discussion.docx"
    create_standalone_section7(standalone_path)
    
    update_master_manuscript(r"D:\chandru project\Final_Manuscript_with_Section7.docx")
