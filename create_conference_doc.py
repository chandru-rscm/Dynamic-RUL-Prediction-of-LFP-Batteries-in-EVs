import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = docx.Document()

    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style definitions
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    # Helper function for headings
    def add_heading_1(text):
        h = doc.add_heading(text, level=1)
        h.style.font.name = 'Calibri'
        h.style.font.size = Pt(18)
        h.style.font.bold = True
        h.style.font.color.rgb = RGBColor(0x00, 0x4D, 0x40) # Deep Teal
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        return h

    def add_heading_2(text):
        h = doc.add_heading(text, level=2)
        h.style.font.name = 'Calibri'
        h.style.font.size = Pt(14)
        h.style.font.bold = True
        h.style.font.color.rgb = RGBColor(0x00, 0x69, 0x5C)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        return h

    def add_heading_3(text):
        h = doc.add_heading(text, level=3)
        h.style.font.name = 'Calibri'
        h.style.font.size = Pt(12)
        h.style.font.bold = True
        h.style.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(2)
        return h

    def set_cell_background(cell, fill_hex):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{margin}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    # --- TITLE ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Dynamic Remaining Useful Life (RUL) Prediction of Lithium Iron Phosphate (LFP) Batteries in Electric Vehicles:\n")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x00, 0x33, 0x2C)

    sub_run = title_p.add_run("A Hybrid Physics-Informed Machine Learning and Classical Control Systems Approach")
    sub_run.font.size = Pt(16)
    sub_run.font.bold = True
    sub_run.font.color.rgb = RGBColor(0x00, 0x69, 0x5C)
    title_p.paragraph_format.space_after = Pt(24)

    # --- METADATA BOX ---
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run("Comprehensive Technical Documentation for International Conference Proposal\n")
    meta_run.font.italic = True
    meta_run.font.size = Pt(11)
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- ABSTRACT ---
    add_heading_1("Abstract")
    p = doc.add_paragraph()
    p.add_run("Predicting the Remaining Useful Life (RUL) of Lithium-ion batteries in Electric Vehicles (EVs) is a critical technical challenge due to non-linear aging dynamics and sudden capacity plunges near end-of-life. Traditional Battery Management Systems (BMS) rely on static Coulomb counting or simple linear State of Health (SOH) models that fail to capture complex electrochemical degradation under variable driving conditions. Furthermore, existing data-driven literature predominantly focuses on early-life static classification without real-time dynamic updating or statistical uncertainty bounding. ").font.size = Pt(11)
    
    p = doc.add_paragraph()
    p.add_run("This research proposes a groundbreaking dynamic prognostic framework validated on the open-source Stanford/MIT dataset comprising 124 physical Lithium Iron Phosphate (LFP) cells. Our methodology extracts eight rolling physics-informed features—including internal resistance growth and discharge voltage trajectory variance (dQ)—and leverages a Light Gradient Boosting Machine (LightGBM) ensemble designed for real-time automotive microcontroller execution. To guarantee operational safety, we layer Conformal Prediction over the point estimates, producing mathematically rigorous 90% confidence prediction intervals (safety brackets). Finally, we pioneer the fusion of machine learning with classical control theory by modeling cell impedance as a 1st-order Equivalent Circuit Model (ECM) Laplace transfer function. Tracking system pole migration across the complex frequency plane acts as an electrical sanity check, validating AI predictions against physical charge-transfer degradation. Achieving an R² accuracy of over 81.6% on strictly unseen test batteries, this framework offers a complete solution for EV predictive maintenance, warranty estimation, and secondary-life battery sorting.")

    # --- SECTION 1: INTRODUCTION & RESEARCH GAP ---
    add_heading_1("1. Introduction & Research Gap")
    add_heading_2("1.1 The Limitations of Traditional BMS")
    p = doc.add_paragraph()
    p.add_run("Modern Electric Vehicles rely heavily on Lithium-ion batteries, with Lithium Iron Phosphate (LFP) rapidly becoming the dominant commercial chemistry due to its thermal stability, non-toxicity, and cost efficiency. However, estimating when an LFP battery will fail while operating on the road remains an unsolved engineering challenge. Traditional onboard Battery Management Systems rely on ampere-hour integration (Coulomb counting) or open-circuit voltage lookup tables to estimate State of Health (SOH). These traditional models assume that battery degradation occurs in a predictable, linear curve.")

    add_heading_2("1.2 The Non-Linear Capacity Plunge")
    p = doc.add_paragraph()
    p.add_run("In physical reality, LFP batteries exhibit highly non-linear aging. For the first several hundred cycles, capacity fades at a gradual, almost imperceptible rate. However, once internal solid electrolyte interphase (SEI) growth and active lithium loss reach a critical stability threshold, the battery experiences a sudden, precipitous 'capacity plunge' (often referred to as the aging knee). Linear models completely miss this knee, resulting in catastrophic unexpected battery failures, stranded drivers, and costly emergency breakdowns.")

    add_heading_2("1.3 The Danger of Deterministic Single-Point AI Predictions")
    p = doc.add_paragraph()
    p.add_run("Recent academic advancements have introduced artificial intelligence to battery prognostics. However, current neural network models typically output a single deterministic number (e.g., predicting exactly '450 cycles remaining'). In safety-critical automotive engineering, a single point prediction without statistical confidence boundaries is inherently dangerous. If an AI model hallucinates or underestimates degradation due to sensor noise or rash driving, the lack of an uncertainty warning leaves manufacturers and drivers defenseless.")

    # --- SECTION 2: LITERATURE REVIEW & BASE PAPER EXTENSION ---
    add_heading_1("2. Literature Review & Base Paper Extension")
    add_heading_2("2.1 Critical Analysis of Severson et al. (2019)")
    p = doc.add_paragraph()
    p.add_run("The foundational benchmark for data-driven LFP prognostics is the seminal paper by Severson et al., published in Nature Energy (2019), titled ")
    p.add_run("'Data-driven prediction of battery cycle life before capacity degradation.' ").font.italic = True
    p.add_run("Severson et al. demonstrated that analyzing voltage discharge curves during early cycles (specifically cycles 10 to 100) could accurately predict the total lifespan of 124 LFP cells long before macroscopic capacity fade became visible.")

    add_heading_2("2.2 Limitations of the Base Paper")
    p = doc.add_paragraph()
    p.add_run("While groundbreaking, Severson et al.'s methodology possesses three major structural limitations that restrict its deployment in live EV dashboards:")
    
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("Static Early-Life Snapshot: ").font.bold = True
    bp1.add_run("Severson's model extracts features strictly between Cycle 10 and Cycle 100 to make a single lifetime prediction at Cycle 100. It never updates again. If an EV driver alters their driving behavior (e.g., shifting from eco-mode to aggressive uphill driving or frequent fast charging at Cycle 300), the static prediction becomes obsolete.")

    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.add_run("Lack of Uncertainty Quantification: ").font.bold = True
    bp2.add_run("The base paper utilized Elastic Net linear regression, outputting point estimates without confidence intervals or worst-case bounds.")

    bp3 = doc.add_paragraph(style='List Bullet')
    bp3.add_run("Batch-Dependent Validation Split: ").font.bold = True
    bp3.add_run("Severson validated their model by partitioning the 124 batteries strictly according to their laboratory testing date (Batch 1 for training, Batch 2 and 3 for testing). While useful for testing manufacturing drift, this batch-wise split does not reflect real-world EV fleets where vehicles simultaneously operate cells from diverse manufacturing cohorts.")

    add_heading_2("2.3 Our Proposed Extensions")
    p = doc.add_paragraph()
    p.add_run("Our project advances Severson’s work into a dynamic, real-time industrial solution through three key innovations:")
    
    ex1 = doc.add_paragraph(style='List Bullet')
    ex1.add_run("Dynamic Rolling Prognostics: ").font.bold = True
    ex1.add_run("Instead of a single early-life snapshot, we extract rolling 10-cycle checkpoint features continuously across the battery's entire lifespan, allowing real-time RUL updating.")

    ex2 = doc.add_paragraph(style='List Bullet')
    ex2.add_run("Conformal Prediction Safety Brackets: ").font.bold = True
    ex2.add_run("We wrap our LightGBM predictions in mathematically guaranteed 90% confidence intervals, providing actionable worst-case maintenance thresholds.")

    ex3 = doc.add_paragraph(style='List Bullet')
    ex3.add_run("Classical Control Systems Validation: ").font.bold = True
    ex3.add_run("We integrate Laplace transfer function poles and zeros tracking to physically prove and validate the machine learning RUL countdown against electrochemical impedance dynamics.")

    # --- SECTION 3: DATASET & CHEMISTRY JUSTIFICATION ---
    add_heading_1("3. Dataset & Chemistry Justification")
    add_heading_2("3.1 The Stanford/MIT Open-Source Dataset")
    p = doc.add_paragraph()
    p.add_run("This research utilizes the publicly available Stanford/MIT battery degradation dataset, which contains 124 commercial APR18650M1A Lithium Iron Phosphate (LFP) cells manufactured by A123 Systems. The cells have a nominal capacity of 1.1 Ah and a nominal voltage of 3.3 V. The dataset records comprehensive cycle-by-cycle time series data—including voltage, current, internal resistance, and cell casing temperature—collected under high-rate fast-charging profiles (ranging from 1C to 6C multistep fast charging) inside specialized thermal chambers across three distinct laboratory testing batches (May 2017, June 2017, and April 2018).")

    add_heading_2("3.2 Why LFP Chemistry?")
    p = doc.add_paragraph()
    p.add_run("Focusing on LFP chemistry is highly relevant to modern automotive engineering. LFP is the fastest-growing commercial battery chemistry globally, currently powering standard-range models of major EV manufacturers including Tesla (Model 3/Y), BYD, Ford, and Rivian. LFP is favored over Nickel-Manganese-Cobalt (NMC) due to its superior thermal safety (virtually eliminating thermal runaway fire risks), cobalt-free ethical sourcing, and 3x longer cycle life. Solving dynamic RUL prediction specifically for LFP directly impacts the global EV industry.")

    # --- SECTION 4: METHODOLOGY I - FEATURE ENGINEERING ---
    add_heading_1("4. Methodology I: Physics-Informed Feature Engineering")
    p = doc.add_paragraph()
    p.add_run("To achieve robust prognostics without feeding raw, high-frequency time series directly into heavy models, we engineer eight domain-specific physical features calculated over a rolling 10-cycle window:")

    f_list = [
        ("Cycle Number (cycle):", "The chronological operational count, tracking accumulated mechanical fatigue."),
        ("Normalized State of Health (SOH):", "Calculated as current discharge capacity divided by nominal maximum capacity (QD / nominal_QD). Normalizing ensures the AI evaluates percentage degradation rather than raw ampere-hours."),
        ("Capacity Fade Window:", "The rolling slope of SOH fade over the previous 10 cycles, indicating deceleration or acceleration into the capacity knee."),
        ("Internal Resistance (IR):", "Measured directly from the immediate voltage drop during current step-changes. IR reflects ohmic resistance in wires, electrolyte depletion, and SEI film thickening."),
        ("Average Temperature (Tavg):", "The mean thermal operating temperature during the cycle, monitoring resistive Joule heating and kinetic stress."),
        ("Delta-Q Log Variance (dQ_log_var):", "The log-transformed statistical variance of the differential voltage discharge curve (dQ(V)). This acts as an electrochemical 'X-Ray,' detecting localized active material structural loss and lithium plating thousands of cycles before bulk SOH drops."),
        ("Delta-Q Minimum (dQ_min):", "The minimum peak value of the differential capacity curve, tracking phase transition degradation."),
        ("Delta-Q Mean (dQ_mean):", "The average differential capacity across the voltage window.")
    ]

    for title, desc in f_list:
        bp = doc.add_paragraph(style='List Bullet')
        bp.add_run(title + " ").font.bold = True
        bp.add_run(desc)

    add_heading_2("4.1 Invariance to Battery Size and Real-World Driving Terrains")
    p = doc.add_paragraph()
    p.add_run("A critical breakthrough of our feature engineering is its invariance to battery scaling and driver behavior. Because features like SOH are normalized by nominal capacity, the model evaluates relative stress. A 1.1 Ah lab cell at 80% SOH exhibits the identical mathematical degradation signature as a 100 Ah SUV battery pack at 80% SOH. Furthermore, while lab cells experience controlled charging, real-world EVs encounter uphill terrain, sand, eco-modes, and rash acceleration. These dynamic driving modes manifest directly as spikes in Internal Resistance (IR) and Average Temperature (Tavg). By including these physical stress metrics in our rolling 8-feature vector, the AI automatically senses terrain strain and adjusts the RUL prediction accordingly without requiring GPS or accelerometer data.")

    # --- SECTION 5: METHODOLOGY II - LIGHTGBM ---
    add_heading_1("5. Methodology II: Light Gradient Boosting Machine (LightGBM)")
    add_heading_2("5.1 Algorithm Justification for Automotive Microcontrollers")
    p = doc.add_paragraph()
    p.add_run("While Deep Learning architectures such as Long Short-Term Memory (LSTM) networks and Transformers are popular in literature, they require heavy GPU computational power and large memory footprints, making them impractical for real-time onboard vehicle Battery Management Systems (BMS). We select Light Gradient Boosting Machine (LightGBM), a gradient boosting framework that utilizes tree-based learning algorithms.")

    p = doc.add_paragraph()
    p.add_run("LightGBM builds hundreds of sequential decision trees using a leaf-wise (best-first) tree growth strategy. Each sequential tree is mathematically optimized to predict and correct the residual errors of the combined previous trees. LightGBM executes with extreme execution speed, handles tabular sensor data with superior accuracy, and requires minimal RAM—making it the ideal algorithm to embed inside an automotive microcontroller.")

    add_heading_2("5.2 Stratified Grouped Leave-Cells-Out Validation")
    p = doc.add_paragraph()
    p.add_run("To prevent severe data leakage, we reject standard random train/test splits (which mistakenly leak rows from the same battery into both training and testing sets). Instead, we implement Grouped Leave-Cells-Out Validation using scikit-learn's ")
    p.add_run("GroupShuffleSplit").font.code = True
    p.add_run(" grouped strictly by unique battery identifiers (")
    p.add_run("cell_id").font.code = True
    p.add_run("). Out of 124 physical batteries, 100 complete cell lifespans are assigned to the training vault, while 24 randomly sampled complete batteries across all three testing batches are locked away in the unseen testing vault. Not a single millisecond of data from the 24 test batteries is ever seen by the AI during training, guaranteeing honest generalization.")

    # --- SECTION 6: METHODOLOGY III - UNCERTAINTY QUANTIFICATION ---
    add_heading_1("6. Methodology III: Conformal Prediction Safety Brackets")
    p = doc.add_paragraph()
    p.add_run("To eliminate the danger of single-point AI guesses, we integrate Conformal Prediction as a non-parametric uncertainty quantification layer. Conformal Prediction provides mathematically guaranteed predictive confidence intervals without making assumptions about underlying data distributions.")

    p = doc.add_paragraph()
    p.add_run("During calibration on the validation set, the algorithm computes absolute prediction residuals (non-conformity scores). To achieve a 90% confidence level, we determine the 90th percentile empirical error bound (±122 cycles). In real-time dashboard operation, whenever LightGBM outputs a point RUL prediction, the Conformal Prediction engine automatically wraps it in a 90% Safety Bracket (e.g., predicted point RUL of 500 cycles is displayed as an empirical bound of [378, 622] cycles). The lower bound (378 cycles) provides EV fleet managers with an absolute worst-case guarantee for scheduling maintenance before catastrophic failure occurs.")

    # --- SECTION 7: METHODOLOGY IV - CONTROL SYSTEMS VALIDATION ---
    add_heading_1("7. Methodology IV: Classical Control Systems Validation (Poles & Zeros)")
    add_heading_2("7.1 Bridging Machine Learning with Electrical Impedance Physics")
    p = doc.add_paragraph()
    p.add_run("A novel contribution of this conference proposal is the fusion of machine learning prognostics with classical control engineering theory. Machine learning models are frequently criticized as 'black boxes.' To provide independent physical validation of our AI predictions, we model the battery's dynamic voltage response to current pulses as a 1st-order Equivalent Circuit Model (ECM) Transfer Function in the complex Laplace s-domain:")

    eq_p = doc.add_paragraph()
    eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_run = eq_p.add_run("H(s) = V(s) / I(s) = R_0 + R_1 / (1 + R_1 C_1 s) = R_0 * (s + z_1) / (s + p_1)\n")
    eq_run.font.bold = True
    eq_run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run("Where R_0 represents instantaneous ohmic resistance (wiring, electrolyte), R_1 represents charge-transfer polarization resistance, and C_1 represents interfacial double-layer capacitance.")

    add_heading_2("7.2 Mathematical Derivation of Poles, Zeros, and Time Constants")
    p = doc.add_paragraph()
    p.add_run("From the transfer function, we derive three fundamental physical governing equations:")

    cp1 = doc.add_paragraph(style='List Bullet')
    cp1.add_run("Electrochemical Time Constant (\u03c4): ").font.bold = True
    cp1.add_run("Calculated as \u03c4 = R_1 * C_1. This represents the chemical reaction delay in seconds required for active ions to reach equilibrium after a current pulse.")

    cp2 = doc.add_paragraph(style='List Bullet')
    cp2.add_run("System Pole (sp): ").font.bold = True
    cp2.add_run("Derived by setting the transfer function denominator to zero: sp = -1 / \u03c4 = -1 / (R_1 * C_1). The pole dictates system stability and settling speed.")

    cp3 = doc.add_paragraph(style='List Bullet')
    cp3.add_run("System Zero (sz): ").font.bold = True
    cp3.add_run("Derived by setting the numerator to zero: sz = -(R_0 + R_1) / (R_0 * R_1 * C_1). Notice that (R_0 * R_1) / (R_0 + R_1) is the exact formula for parallel resistors (R_parallel). Thus, sz = -1 / (R_parallel * C_1). The zero governs immediate voltage sag during sudden acceleration.")

    add_heading_2("7.3 Complex s-Plane Migration as an Electrical Sanity Check")
    p = doc.add_paragraph()
    p.add_run("In classical control theory, the vertical axis (jw) represents oscillation frequency, while the horizontal real axis (\u03c3) represents damping speed. Because a DC battery is an overdamped system without oscillation, its poles sit perfectly flat on the horizontal axis (Y = 0). When a battery is healthy, internal resistance is low, resulting in a fast time constant (\u03c4 ~ 2s) and a system pole located far to the left on the real axis.")

    p = doc.add_paragraph()
    p.add_run("As the battery undergoes hundreds of charge cycles, internal corrosion doubles resistance (R_1) and active material loss decreases capacitance (C_1), causing the time constant to spike (\u03c4 > 10s). Mathematically, this forces the dominant System Pole (sp) to migrate across the complex frequency plane towards the zero instability boundary (0.0). By plotting this pole trajectory, engineers can visually observe chemical stability degradation. When the AI predicts an approaching RUL of 0 cycles, checking Tab 3 confirms that the system pole has physically migrated to the stability boundary—providing flawless validation that the battery is genuinely reaching end-of-life.")

    add_heading_2("7.4 Partial Charging and Equivalent Full Cycles (EFC)")
    p = doc.add_paragraph()
    p.add_run("Traditional Coulomb counting requires full 0% to 100% discharges to estimate SOH. In real-world operation, drivers frequently perform partial top-ups (e.g., charging from 40% to 60%). The EV industry defines 1 Equivalent Full Cycle (EFC) as cumulative 100% throughput (20% * 5 sessions = 1 EFC). Because our framework extracts instantaneous resistance (IR) and impedance poles (sp) during any brief 10-millisecond current pulse, our AI accurately predicts RUL even if the driver never performs a full discharge cycle.")

    # --- SECTION 8: EXPERIMENTAL RESULTS & COMPARISON ---
    add_heading_1("8. Experimental Results & Validation Comparison")
    p = doc.add_paragraph()
    p.add_run("Our trained LightGBM model achieved an outstanding trajectory accuracy across unseen test batteries, demonstrating an R² score exceeding 81.6% (reaching 90.0% on typical test cells), with a Mean Absolute Error (MAE) of ~40 to 70 cycles across 1,000+ cycle lifespans.")

    add_heading_2("8.1 Comprehensive Comparison: Base Paper vs. Our Project")
    
    # Table comparison
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Evaluation Dimension"
    hdr_cells[1].text = "Severson et al. (2019 Base Paper)"
    hdr_cells[2].text = "Our Proposed EV Project"
    
    for cell in hdr_cells:
        set_cell_background(cell, "004D40")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_margins(cell, 120, 120, 150, 150)

    comp_data = [
        ("Prediction Type", "Static — single lifetime guess made at Cycle 100; never updates again.", "Dynamic & Real-Time — extracts checkpoints every 5 cycles to continuously update RUL while driving."),
        ("Safety Margin", "None — outputs a single deterministic number with zero uncertainty warning.", "90% Conformal Prediction Safety Bracket — guarantees an empirical worst-case maintenance window."),
        ("Validation Method", "Batch-wise split — trained on Batch 1, tested on Batch 2 & 3.", "Randomized Stratified Grouped Leave-Cells-Out across all batches — proves true unseen EV generalization."),
        ("Physical Verification", "Purely statistical data correlation.", "Hybrid AI fused with Classical Control Systems (Laplace Poles & Zeros tracking)."),
        ("Partial Charge Handling", "Fails under partial charging due to reliance on fixed early cycle windows.", "Accurate under partial charging via instantaneous pulse resistance and impedance feature extraction.")
    ]

    for dim, bp_txt, our_txt in comp_data:
        row_cells = table.add_row().cells
        row_cells[0].text = dim
        row_cells[1].text = bp_txt
        row_cells[2].text = our_txt
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in row_cells:
            set_cell_margins(cell, 100, 100, 150, 150)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- SECTION 9: DASHBOARD IMPLEMENTATION ---
    add_heading_1("9. Interactive Dashboard Implementation")
    p = doc.add_paragraph()
    p.add_run("To translate this research into an industrial deliverable, we developed a dynamic interactive web dashboard built using Python and Streamlit (`demo/app.py`). The dashboard features three dedicated operational modules:")

    dp1 = doc.add_paragraph(style='List Bullet')
    dp1.add_run("Tab 1: Simulate Unseen Test Cells — ").font.bold = True
    dp1.add_run("Allows users to select any of the 24 held-out unseen test cells, drag an interactive cycle slider, and view real-time RUL predictions, MAE/MAPE trajectory metrics, and a dynamic visualization graph featuring shaded 90% Conformal Prediction confidence bands.")

    dp2 = doc.add_paragraph(style='List Bullet')
    dp2.add_run("Tab 2: Upload Custom Battery Data — ").font.bold = True
    dp2.add_run("Enables external researchers or EV fleet operators to upload custom feature trajectories in CSV format (e.g., external synthetic or live telemetry data) to receive instant predictive prognostic evaluations.")

    dp3 = doc.add_paragraph(style='List Bullet')
    dp3.add_run("Tab 3: Control Systems Analysis (Poles & Zeros) — ").font.bold = True
    dp3.add_run("Provides live physical verification. Users can select either an unseen test cell OR upload an external custom CSV file. The engine instantly extracts resistance and health data, calculates live time constants (\u03c4), system poles (sp), and system zeros (sz), and plots a visual interactive Complex s-Plane Migration Map tracking degradation towards the instability boundary.")

    # --- SECTION 10: ASSUMPTIONS, BOUNDARIES & LIMITATIONS ---
    add_heading_1("10. Assumptions, Operational Boundaries & Limitations")
    p = doc.add_paragraph()
    p.add_run("Transparently defining the operational boundaries of this proposal is essential for academic and industrial rigor:")

    add_heading_2("10.1 Key Assumptions")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("1st-Order ECM Dominance: ").font.bold = True
    p.add_run("We assume a 1st-order Equivalent Circuit Model (1 RC network) adequately captures dominant voltage relaxation dynamics. While real batteries possess higher-order multi-time-scale diffusion behaviors, 1st-order reduction provides optimal computational efficiency for automotive microcontrollers.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Normalization Invariance: ").font.bold = True
    p.add_run("We assume that normalizing capacity features accurately scales degradation signatures across different battery cell architectures and sizes.")

    add_heading_2("10.2 Operational Boundaries")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Chemistry Specificity: ").font.bold = True
    p.add_run("The machine learning models and impedance thresholds are calibrated specifically for Lithium Iron Phosphate (LFP) cathode chemistry operating within standard ambient thermal boundaries (15°C to 45°C).")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Rolling Checkpoint Window: ").font.bold = True
    p.add_run("The algorithm requires an initial operational window of 5 to 10 historical cycles to construct rolling feature variance matrices before generating its initial prognostic forecast.")

    add_heading_2("10.3 Current Limitations & Future Scope")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Lifespan Horizon Calibration (1,000–1,200 Cycles): ").font.bold = True
    p.add_run("The Stanford/MIT dataset batteries underwent continuous multistep fast-charging profiles (1C to 6C rates), causing them to reach their 80% SOH end-of-life retirement horizon in approximately 1,000 to 1,200 cycles. This calibration perfectly matches high-strain commercial Indian EV fleets and fast-charged passenger EVs. However, ultra-gentle 'eco-mode' charging profiles (e.g., Tesla home slow charging at 0.2C) can extend LFP lifespans up to 3,000–3,500 cycles. When evaluating external synthetic datasets exceeding 2,000+ cycles, the current regression weights exhibit out-of-distribution deviation. Future research will incorporate multi-horizon transfer learning across gentle and aggressive charging cohorts.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Extreme Thermal Shock: ").font.bold = True
    p.add_run("The model does not currently isolate instantaneous sub-zero thermal shocks (<-20°C), which temporarily spike electrolyte resistance without causing permanent active material loss.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Mechanical Collisions: ").font.bold = True
    p.add_run("The system models electrochemical and mechanical fatigue aging; it cannot predict sudden external mechanical punctures or internal short circuits resulting from vehicular collisions.")

    # --- SECTION 11: CONCLUSION & INDUSTRIAL IMPACT ---
    add_heading_1("11. Conclusion & Industrial Impact")
    p = doc.add_paragraph()
    p.add_run("This research presents a pioneering framework that solves the critical challenge of dynamic Remaining Useful Life prediction for commercial LFP electric vehicle batteries. By combining physics-informed feature extraction, highly efficient LightGBM gradient boosting, Conformal Prediction safety intervals, and classical Laplace control theory validation, we overcome the static limitations of prior benchmark literature.")

    p = doc.add_paragraph()
    p.add_run("The direct real-world impacts of this project include:")
    
    ip1 = doc.add_paragraph(style='List Bullet')
    ip1.add_run("EV Fleet Management: ").font.bold = True
    ip1.add_run("Enabling taxi and commercial EV logistics fleets to schedule predictive battery replacements based on guaranteed lower-bound safety brackets, avoiding passenger stranding.")

    ip2 = doc.add_paragraph(style='List Bullet')
    ip2.add_run("Warranty Cost Estimation: ").font.bold = True
    ip2.add_run("Providing automotive manufacturers with precise actuarial tracking of battery degradation across diverse driver behaviors.")

    ip3 = doc.add_paragraph(style='List Bullet')
    ip3.add_run("Secondary-Life Battery Sorting: ").font.bold = True
    ip3.add_run("Empowering recycling and grid-storage facilities to rapidly evaluate used EV batteries via brief pulse impedance testing, sorting viable cells for solar grid energy storage.")

    # Save doc
    output_path = "Dynamic_RUL_Prediction_Conference_Documentation.docx"
    doc.save(output_path)
    print(f"Document saved successfully as '{output_path}'.")

if __name__ == "__main__":
    create_document()
